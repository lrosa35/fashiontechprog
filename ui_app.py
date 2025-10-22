import os
from typing import Optional

import os
from datetime import datetime
from typing import Optional

from fastapi import FastAPI, Request, Form, status, Depends, HTTPException, UploadFile, File
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import ValidationError
from fastapi.security import HTTPBasic, HTTPBasicCredentials
import secrets

# Reuse existing API app and logic
from server import app as api_app
from server import OrcamentoIn, criar_orcamento, listar_orcamentos, obter_orcamento
import orcamento as orc
from db_backend import DB
from openpyxl import load_workbook
import tempfile


app = FastAPI(title="Orçamentos Web UI")

# Mount the JSON API under /api (no changes required in server.py)
app.mount("/api", api_app)


BASE_DIR = os.path.dirname(__file__)
TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")
STATIC_DIR = os.path.join(BASE_DIR, "static")

templates = Jinja2Templates(directory=TEMPLATES_DIR)

if os.path.isdir(STATIC_DIR):
    app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

# Basic auth (optional). Keep above route definitions.
security = HTTPBasic(auto_error=False)
UI_USER = os.getenv("UI_BASIC_USER")
UI_PASS = os.getenv("UI_BASIC_PASS")
UI_DISABLE_AUTH = os.getenv("UI_DISABLE_AUTH", "0") == "1"

def require_auth(credentials: Optional[HTTPBasicCredentials] = Depends(security)):
    if UI_DISABLE_AUTH or not (UI_USER and UI_PASS):
        return  # auth disabled in dev or when not configured
    if credentials is None:
        raise HTTPException(status_code=401, detail="Not authenticated")
    is_user = secrets.compare_digest(credentials.username, UI_USER)
    is_pass = secrets.compare_digest(credentials.password, UI_PASS)
    if not (is_user and is_pass):
        raise HTTPException(status_code=401, detail="Unauthorized")
    return


@app.get("/", response_class=HTMLResponse)
async def index(request: Request, id: Optional[str] = None, cnpj: Optional[str] = None, _auth=Depends(require_auth)):
    rows = None
    if id or cnpj:
        try:
            resp = await listar_orcamentos(id=id, cnpj=cnpj, vendedor=None, start=None, end=None)
            rows = resp.get('rows', []) if isinstance(resp, dict) else []
        except Exception:
            rows = []
    return templates.TemplateResponse(
        'index.html',
        {
            'request': request,
            'defaults': {
                'tipo_servico': 'Impressão',
                'status': 'Sem desconto',
                'unidade': 'Centímetros',
            },
            'rows': rows,
            'q_id': id or '',
            'q_cnpj': cnpj or '',
        },
    )

@app.get("/clientes", response_class=HTMLResponse)
async def clientes(request: Request, doc: Optional[str] = None, _auth=Depends(require_auth)):
    cad = None
    error = None
    if doc:
        try:
            cad = DB.buscar_cadastro_por_documento('CNPJ/CPF', doc)
        except Exception as ex:
            error = f'Falha ao buscar: {ex}'
    return templates.TemplateResponse('clientes.html', {'request': request, 'doc': doc or '', 'cad': cad, 'error': error})

@app.get("/contrato", response_class=HTMLResponse)
async def contrato(request: Request, _auth=Depends(require_auth)):
    return templates.TemplateResponse("contrato.html", {"request": request})

@app.get("/relatorios", response_class=HTMLResponse)
async def relatorios(request: Request, _auth=Depends(require_auth)):
    return templates.TemplateResponse("relatorios.html", {"request": request})

@app.get("/usuarios", response_class=HTMLResponse)
async def usuarios(request: Request, _auth=Depends(require_auth)):
    return templates.TemplateResponse("usuarios.html", {"request": request})


@app.post("/orcamentos", response_class=HTMLResponse)
async def create_orcamento(
    request: Request,
    tipo_servico: str = Form(...),
    cliente: str = Form(...),
    cnpj: str = Form(...),
    email: str = Form(...),
    status_: str = Form(alias="status", default="Sem desconto"),
    unidade: str = Form(...),
    quantidade: str = Form(...),
    vendedor: Optional[str] = Form(default=None),
    forma_pagamento: Optional[str] = Form(default=None),
    preco_por_metro_opc: Optional[str] = Form(default=None),
    metros_opc: Optional[str] = Form(default=None),
    _auth=Depends(require_auth),
):
    try:
        body = OrcamentoIn(
            tipo_servico=tipo_servico,
            cliente=cliente,
            cnpj=cnpj,
            email=email,
            status=status_,
            unidade=unidade,
            quantidade=quantidade,
            vendedor=vendedor,
            forma_pagamento=forma_pagamento,
            preco_por_metro_opc=preco_por_metro_opc,
            metros_opc=metros_opc,
        )
    except ValidationError as ve:
        return templates.TemplateResponse(
            "index.html",
            {
                "request": request,
                "error": "Erro de validação: verifique os campos.",
                "details": ve.errors(),
                "defaults": {
                    "tipo_servico": tipo_servico,
                    "cliente": cliente,
                    "cnpj": cnpj,
                    "email": email,
                    "status": status_,
                    "unidade": unidade,
                    "quantidade": quantidade,
                },
            },
            status_code=status.HTTP_400_BAD_REQUEST,
        )

    try:
        result = await criar_orcamento(body)  # reuse API business logic
        return templates.TemplateResponse(
            "result.html",
            {"request": request, "orc": result.model_dump()},
            status_code=status.HTTP_201_CREATED,
        )
    except Exception as ex:
        return templates.TemplateResponse(
            "index.html",
            {
                "request": request,
                "error": f"Falha ao criar orçamento: {ex}",
                "defaults": body.model_dump(),
            },
            status_code=status.HTTP_400_BAD_REQUEST,
        )


@app.get("/buscar", response_class=HTMLResponse)
async def buscar_page(request: Request, id: Optional[str] = None, cnpj: Optional[str] = None, _auth=Depends(require_auth)):
    rows = []
    error = None
    if id or cnpj:
        try:
            resp = await listar_orcamentos(id=id, cnpj=cnpj, vendedor=None, start=None, end=None)
            rows = resp.get("rows", []) if isinstance(resp, dict) else []
        except Exception as ex:
            error = f"Falha ao buscar: {ex}"
    return templates.TemplateResponse(
        "list.html",
        {"request": request, "rows": rows, "id": id or "", "cnpj": cnpj or "", "error": error},
    )


@app.get("/orcamentos/{orc_id}", response_class=HTMLResponse)
async def detalhe_orcamento(request: Request, orc_id: str, _auth=Depends(require_auth)):
    try:
        d = await obter_orcamento(orc_id)
        return templates.TemplateResponse("detail.html", {"request": request, "orc": d})
    except Exception as ex:
        return templates.TemplateResponse(
            "detail.html",
            {"request": request, "error": f"Não foi possível obter o orçamento: {ex}"},
            status_code=status.HTTP_404_NOT_FOUND,
        )


def _infer_doc_label_and_value(d: dict) -> tuple[str, str, str]:
    # returns (documento, label, value)
    doc_raw = d.get("CNPJ/CPF") or d.get("cnpj") or d.get("cnpj_cpf") or ""
    digits = "".join([ch for ch in str(doc_raw) if ch.isdigit()])
    if len(digits) == 11:
        return ("CPF", "Nome", orc.formatar_cpf(digits))
    if len(digits) == 14:
        return ("CNPJ", "RazÃ£o Social", orc.formatar_cnpj(digits))
    # fallback
    return ("Documento", "Documento", str(doc_raw))


@app.get("/orcamentos/{orc_id}/pdf")
async def baixar_pdf(orc_id: str, _auth=Depends(require_auth)):
    # get data
    d = await obter_orcamento(orc_id)
    if not isinstance(d, dict):
        raise HTTPException(404, "OrÃ§amento nÃ£o encontrado")

    id_orc = d.get("ID OrÃ§amento") or d.get("id_orcamento") or orc_id
    datahora = d.get("Data/Hora") or d.get("data_hora") or datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    tipo_servico = d.get("Tipo de ServiÃ§o") or d.get("tipo_servico") or "Impressão"
    cliente_val = d.get("CLIENTE (Valor)") or d.get("Cliente") or d.get("cliente") or ""
    vendedor = d.get("Vendedor") or ""
    status_val = d.get("Status") or d.get("status") or "Sem desconto"
    qtd = d.get("Quantidade") or d.get("quantidade") or ""
    unidade = d.get("Unidade") or d.get("unidade") or "Centímetros"
    metros = d.get("Metros") or d.get("metros") or ""
    preco = d.get("PreÃ§o por metro") or d.get("preco_por_metro") or ""
    forma_pgto = d.get("Forma de Pagamento") or ""
    total = d.get("Valor Total") or d.get("valor_total") or ""

    documento, cliente_label, doc_fmt = _infer_doc_label_and_value(d)

    dados_seq = [
        id_orc,
        datahora,
        tipo_servico,
        cliente_label,
        str(cliente_val),
        documento,
        doc_fmt,
        str(d.get("E-mail") or d.get("email") or ""),
        vendedor,
        status_val,
        str(qtd),
        unidade,
        str(metros),
        str(preco),
        str(forma_pgto),
        str(total),
    ]

    export_dir = os.path.join(os.path.dirname(__file__), "data", "exports")
    os.makedirs(export_dir, exist_ok=True)
    pdf_path = orc.gerar_pdf_orcamento(dados_seq, export_dir)
    filename = os.path.basename(pdf_path)
    return FileResponse(pdf_path, media_type="application/pdf", filename=filename)


@app.get("/healthz")
async def healthz():
    return {"ok": True}




# ===== Importar Planilha (Excel -> DB) =====
def _header_map(ws):
    return [(c.value if c is not None else None) for c in ws[1]]


def _row_to_dict(headers, row):
    out = {}
    for i, h in enumerate(headers):
        if not h:
            continue
        v = row[i].value if i < len(row) else None
        out[str(h)] = v if v is not None else ""
    return out


@app.get('/importar', response_class=HTMLResponse)
async def importar_get(request: Request, _auth=Depends(require_auth)):
    return templates.TemplateResponse('importar.html', {'request': request})


@app.post('/importar', response_class=HTMLResponse)
async def importar_post(request: Request, file: UploadFile = File(...), _auth=Depends(require_auth)):
    if not DB.is_ready():
        return templates.TemplateResponse('importar.html', {'request': request, 'error': 'Banco não configurado. No Heroku, adicione o add-on Postgres (DATABASE_URL é criado automaticamente).'}, status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name
    wb = load_workbook(tmp_path, read_only=True, data_only=True)
    def get_ws(*names):
        for n in names:
            if n in wb.sheetnames:
                return wb[n]
        return None
    ws_orc = get_ws('Orçamentos', 'Orcamentos')
    ws_cad = get_ws('Cadastros')
    ws_ped = get_ws('Pedidos')
    counts = {'orc_lidos': 0, 'orc_inseridos': 0, 'cad_lidos': 0, 'cad_inseridos': 0, 'ped_lidos': 0, 'ped_inseridos': 0}
    if ws_orc is not None:
        headers = _header_map(ws_orc)
        for row in ws_orc.iter_rows(min_row=2):
            d = _row_to_dict(headers, row)
            if not any(v for v in d.values()):
                continue
            counts['orc_lidos'] += 1
            try:
                DB.salvar_orcamento(d)
                counts['orc_inseridos'] += 1
            except Exception:
                pass
    if ws_cad is not None:
        headers = _header_map(ws_cad)
        for row in ws_cad.iter_rows(min_row=2):
            d = _row_to_dict(headers, row)
            if not any(v for v in d.values()):
                continue
            counts['cad_lidos'] += 1
            try:
                DB.salvar_cadastro(d)
                counts['cad_inseridos'] += 1
            except Exception:
                pass
    if ws_ped is not None:
        headers = _header_map(ws_ped)
        for row in ws_ped.iter_rows(min_row=2):
            d = _row_to_dict(headers, row)
            if not any(v for v in d.values()):
                continue
            counts['ped_lidos'] += 1
            try:
                DB.salvar_pedido(d)
                counts['ped_inseridos'] += 1
            except Exception:
                pass
    return templates.TemplateResponse('importar.html', {'request': request, 'result': counts})




@app.post('/clientes', response_class=HTMLResponse)
async def clientes_post(request: Request, _auth=Depends(require_auth)):
    form = await request.form()
    d = {k: str(v) for k, v in form.items()}
    if 'cnpj_cpf' in d and 'CNPJ/CPF' not in d:
        d['CNPJ/CPF'] = d['cnpj_cpf']
    try:
        DB.salvar_cadastro(d)
        msg = 'Cadastro salvo com sucesso.'
        cad = DB.buscar_cadastro_por_documento('CNPJ/CPF', d.get('CNPJ/CPF',''))
    except Exception as ex:
        msg = f'Falha ao salvar: {ex}'
        cad = d
    return templates.TemplateResponse('clientes.html', {'request': request, 'doc': d.get('CNPJ/CPF',''), 'cad': cad, 'error': None, 'msg': msg})
@app.post('/contrato', response_class=HTMLResponse)
async def contrato_post(request: Request, _auth=Depends(require_auth)):
    form = await request.form()
    orc_id = str(form.get('orc_id') or '').strip()
    if not orc_id:
        return templates.TemplateResponse('contrato.html', {'request': request, 'error': 'Informe o ID do orçamento'}, status_code=status.HTTP_400_BAD_REQUEST)
    try:
        d = await obter_orcamento(orc_id)
        if not isinstance(d, dict):
            raise RuntimeError('Orçamento não encontrado')
    except Exception as ex:
        return templates.TemplateResponse('contrato.html', {'request': request, 'error': f'Falha ao obter orçamento: {ex}', 'orc_id': orc_id}, status_code=status.HTTP_404_NOT_FOUND)

    export_dir = os.path.join(os.path.dirname(__file__), 'data', 'exports')
    os.makedirs(export_dir, exist_ok=True)
    out_path = os.path.join(export_dir, f'Contrato_{orc_id}.docx')

    try:
        from docx import Document
        tdir = os.path.join(os.path.dirname(__file__), 'data', 'CONTRATO PARA ATUALIZAÇÃO')
        template = None
        if os.path.isdir(tdir):
            for name in os.listdir(tdir):
                if name.lower().endswith('.docx'):
                    template = os.path.join(tdir, name); break
        doc = Document(template) if template else Document()
        if not template:
            doc.add_heading('Contrato - Orçamento ' + orc_id, 0)
        kv = {
            'ID Orçamento': orc_id,
            'Data/Hora': d.get('Data/Hora') or d.get('data_hora'),
            'Tipo de Serviço': d.get('Tipo de Serviço') or d.get('tipo_servico'),
            'Cliente': d.get('CLIENTE (Valor)') or d.get('Cliente') or d.get('cliente'),
            'CNPJ/CPF': d.get('CNPJ/CPF') or d.get('cnpj') or d.get('cnpj_cpf'),
            'E-mail': d.get('E-mail') or d.get('email'),
            'Vendedor': d.get('Vendedor') or '',
            'Status': d.get('Status') or d.get('status') or '',
            'Quantidade': d.get('Quantidade') or d.get('quantidade'),
            'Unidade': d.get('Unidade') or d.get('unidade'),
            'Metros': d.get('Metros') or d.get('metros'),
            'Preço por metro': d.get('Preço por metro') or d.get('preco_por_metro'),
            'Forma de Pagamento': d.get('Forma de Pagamento') or '',
            'Valor Total': d.get('Valor Total') or d.get('valor_total'),
        }
        for p in doc.paragraphs:
            for k,v in kv.items():
                if v is None: v = ''
                p.text = p.text.replace('{{'+k+'}}', str(v))
        if not template:
            table = doc.add_table(rows=0, cols=2)
            for k,v in kv.items():
                row = table.add_row().cells
                row[0].text = k
                row[1].text = str(v or '')
        doc.save(out_path)
    except Exception as ex:
        return templates.TemplateResponse('contrato.html', {'request': request, 'error': f'Falha ao gerar DOCX: {ex}', 'orc_id': orc_id}, status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)

    dl_url = f"/static-download?path={out_path}"
    return templates.TemplateResponse('contrato.html', {'request': request, 'download_url': dl_url, 'orc_id': orc_id})

@app.get('/static-download')
async def static_download(path: str, _auth=Depends(require_auth)):
    return FileResponse(path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=os.path.basename(path))
