# -*- coding: utf-8 -*-
import os
import re
import urllib.parse
import sys
from datetime import datetime, timedelta

import flet as ft
import zipfile
import shutil
from openpyxl import Workbook, load_workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

# =========================================================
#                     FONTES / ESTILO
# =========================================================
def register_arial():
    """Tenta registrar Arial do Windows; se falhar, usa Helvetica padrão do ReportLab."""
    try:
        pdfmetrics.registerFont(TTFont("Arial", r"C:\Windows\Fonts\arial.ttf"))
        pdfmetrics.registerFont(TTFont("Arial-Bold", r"C:\Windows\Fonts\arialbd.ttf"))
        pdfmetrics.registerFont(TTFont("Arial-Italic", r"C:\Windows\Fonts\ariali.ttf"))
        return True
    except Exception:
        return False


ARIAL_OK = register_arial()
FONT_REG = "Arial" if ARIAL_OK else "Helvetica"
FONT_BOLD = "Arial-Bold" if ARIAL_OK else "Helvetica-Bold"
FONT_ITAL = "Arial-Italic" if ARIAL_OK else "Helvetica-Oblique"

# Debug opcional (para rodar via VS Code): defina ORC_DEBUG=1
ORC_DEBUG = os.environ.get("ORC_DEBUG", "0") == "1"

MM = 2.834645669
MARGEM = 20 * MM
LINE = 18
H1_SIZE = 22
BODY_SIZE = 12
TOTAL_SIZE = 16
AZUL = colors.HexColor("#1272EB")

HEADER_LEFT_OFFSET = -8
HEADER_CENTER_OFFSET = +16
TITLE_TO_BODY_GAP = LINE * 2.4

LOGO_WIDTH = 110
LOGO_HEIGHT = 88
LOGO_TOP_MARGIN = 1.5 * MM
LOGO_RIGHT_MARGIN = 5 * MM

# =========================================================
#                         HELPERS
# =========================================================
def http_get_json(url: str, timeout: int = 8):
    """Tenta 'requests'; se não houver, usa 'urllib'."""
    try:
        import requests  # type: ignore
        r = requests.get(url, timeout=timeout)
        r.raise_for_status()
        return r.json()
    except Exception:
        import urllib.request
        import json as _json
        with urllib.request.urlopen(url, timeout=timeout) as resp:
            return _json.loads(resp.read().decode("utf-8"))


def http_post_json(url: str, payload: dict, timeout: int = 10):
    """POST JSON com fallback para urllib."""
    try:
        import requests  # type: ignore
        r = requests.post(url, json=payload, timeout=timeout)
        r.raise_for_status()
        return r.json()
    except Exception:
        import urllib.request
        import json as _json
        data = _json.dumps(payload).encode("utf-8")
        req = urllib.request.Request(url, data=data, headers={"Content-Type": "application/json"}, method="POST")
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return _json.loads(resp.read().decode("utf-8"))


def _read_text_file(path: str) -> str | None:
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except Exception:
        return None


def get_api_base() -> str:
    # Procura cloud_api_url.txt em locais comuns ao rodar via PyInstaller
    candidates = []
    try:
        if getattr(sys, "frozen", False) and hasattr(sys, "executable"):
            candidates.append(os.path.dirname(sys.executable))
    except Exception:
        pass
    try:
        candidates.append(os.path.dirname(os.path.abspath(__file__)))
    except Exception:
        pass
    candidates.append(os.getcwd())

    for base in candidates:
        p = _read_text_file(os.path.join(base, "cloud_api_url.txt"))
        if p:
            base_url = p.rstrip("/")
            if ORC_DEBUG:
                print(f"[DEBUG] API base via cloud_api_url.txt: {base_url}")
            return base_url

    env = os.environ.get("CLOUD_API_DEFAULT_URL")
    if env:
        base_url = env.rstrip("/")
        if ORC_DEBUG:
            print(f"[DEBUG] API base via env: {base_url}")
        return base_url

    if ORC_DEBUG:
        print("[DEBUG] API base default: http://localhost:8000")
    return "http://localhost:8000"


def api_get(path: str) -> dict:
    import urllib.parse as _up
    base = get_api_base()
    if not path.startswith("/"):
        path = "/" + path
    url = base + path
    if ORC_DEBUG:
        print(f"[DEBUG] GET {url}")
    return http_get_json(url)


def api_post(path: str, payload: dict) -> dict:
    base = get_api_base()
    if not path.startswith("/"):
        path = "/" + path
    url = base + path
    if ORC_DEBUG:
        print(f"[DEBUG] POST {url} payload_keys={list(payload.keys())}")
    return http_post_json(url, payload)


def find_onedrive_base() -> str | None:
    """Descobre a pasta base do OneDrive (qualquer conta)."""
    for env in ["OneDriveCommercial", "OneDrive", "OneDriveConsumer"]:
        p = os.environ.get(env)
        if p and os.path.isdir(p):
            return p
    home = os.path.expanduser("~")
    try:
        for nome in os.listdir(home):
            if nome.lower().startswith("onedrive"):
                p = os.path.join(home, nome)
                if os.path.isdir(p):
                    return p
    except Exception:
        pass
    return None


def safe_join(*parts: str) -> str:
    return os.path.join(*[str(x) for x in parts if x is not None])


def try_first_existing(paths: list[str]) -> str | None:
    for p in paths:
        if p and os.path.exists(p):
            return p
    return None


def sanitize_filename(name: str, replacement: str = "_") -> str:
    """Remove caracteres inválidos para nomes de arquivo no Windows.
    Mantém letras, Números, espaço, hífen, sublinhado, parênteses e ponto.
    """
    import re as _re
    n = (name or "").strip()
    # Substitui caracteres proibidos: <>:"/\|?*
    n = _re.sub(r'[<>:"/\\|?*]', replacement, n)
    # Normaliza espaços
    n = _re.sub(r"\s+", " ", n)
    # Evita nomes muito longos
    return n[:150]


def extrair_nome_CLIENTE(d: dict) -> str:
    """Obtém o nome do CLIENTE, priorizando 'Razão Social/Nome'."""
    return str(
        d.get("Razão Social/Nome")
        or d.get("CLIENTE (Valor)")
        or d.get("CLIENTE")
        or d.get("CLIENTE (Etiqueta PDF)")
        or ""
    )


def find_in_folder(root: str, filename_contains: str, exts: tuple[str, ...]) -> str | None:
    """Busca recursiva por um arquivo que contenha 'filename_contains' e termine com uma extensão do tuple."""
    if not (root and os.path.isdir(root)):
        return None
    lc = filename_contains.lower()
    for base, _, files in os.walk(root):
        for f in files:
            if f.lower().endswith(exts) and lc in f.lower():
                return os.path.join(base, f)
    return None


def validar_cnpj(cnpj: str) -> bool:
    n = re.sub(r"\D", "", cnpj or "")
    if len(n) != 14 or n == n[0] * 14:
        return False
    p1 = [5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]
    p2 = [6] + p1
    r1 = sum(int(n[i]) * p1[i] for i in range(12)) % 11
    dv1 = 0 if r1 < 2 else 11 - r1
    r2 = sum(int(n[i]) * p2[i] for i in range(13)) % 11
    dv2 = 0 if r2 < 2 else 11 - r2
    return n[-2:] == f"{dv1}{dv2}"


def formatar_cnpj(cnpj: str) -> str:
    n = re.sub(r"\D", "", cnpj or "")[:14]
    return f"{n[:2]}.{n[2:5]}.{n[5:8]}/{n[8:12]}-{n[12:]}" if len(n) == 14 else cnpj


def validar_cpf(cpf: str, only_rj: bool = True) -> bool:
    n = re.sub(r"\D", "", cpf or "")
    if len(n) != 11 or n == n[0] * 11:
        return False
    if only_rj and n[8] != "7":
        return False  # CPF RJ
    s1 = sum(int(n[i]) * (10 - i) for i in range(9))
    dv1 = (s1 * 10) % 11
    dv1 = 0 if dv1 == 10 else dv1
    s2 = sum(int(n[i]) * (11 - i) for i in range(10))
    dv2 = (s2 * 10) % 11
    dv2 = 0 if dv2 == 10 else dv2
    return n[-2:] == f"{dv1}{dv2}"


def formatar_cpf(cpf: str) -> str:
    n = re.sub(r"\D", "", cpf or "")[:11]
    return f"{n[:3]}.{n[3:6]}.{n[6:9]}-{n[9:]}" if len(n) == 11 else cpf


def validar_email(email: str) -> bool:
    return re.match(r"^[\w\.-]+@[\w\.-]+\.[a-zA-Z]{2,}(?:\.[a-zA-Z]{2,})*$", email or "") is not None


def formatar_cep(cep: str) -> str:
    n = re.sub(r"\D", "", cep or "")[:8]
    return f"{n[:5]}-{n[5:]}" if len(n) == 8 else cep


def format_num_ptbr(n: float) -> str:
    return f"{n:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _parse_ptbr_float(txt: str) -> float:
    txt = (txt or "").strip()
    try:
        return float(txt.replace(".", "").replace(",", "."))
    except Exception:
        return 0.0


def data_hora_tokens(d: datetime | None = None):
    d = d or datetime.now()
    return {"data_compacta": d.strftime("%d%m%Y"), "combinado": d.strftime("%d/%m/%Y %H:%M:%S")}


def formatar_doc(tipo: str, valor: str) -> str:
    return formatar_cnpj(valor) if tipo == "CNPJ" else formatar_cpf(valor) if tipo == "CPF" else valor


def validar_doc(tipo: str, valor: str) -> bool:
    return validar_cnpj(valor) if tipo == "CNPJ" else validar_cpf(valor, True) if tipo == "CPF" else False


def montar_endereco_entrega_formatado(cad: dict) -> str:
    """LOGRADOURO, Número, COMPLEMENTO, BAIRRO, Município/UF, CEP: 00000-000"""
    log = str(cad.get("Entrega Endereço") or cad.get("Endereço") or "").strip()
    num = str(cad.get("Entrega Número") or cad.get("Número") or "").strip()
    comp = str(cad.get("Entrega Complemento") or cad.get("Complemento") or "").strip()
    bai = str(cad.get("Entrega Bairro") or cad.get("Bairro") or "").strip()
    mun = str(cad.get("Entrega Município") or cad.get("Município") or "").strip()
    uf = str(cad.get("Entrega UF") or cad.get("UF") or "").strip()
    cep_raw = str(cad.get("Entrega CEP") or cad.get("CEP") or "").strip()
    cep_fmt = formatar_cep(cep_raw) if cep_raw else ""

    partes = []
    if log:
        partes.append(log)
    if num:
        partes.append(num)
    if comp:
        partes.append(comp)
    if bai:
        partes.append(bai)

    if mun or uf:
        munuf = (mun or "") + (f"/{uf}" if uf else "")
        if munuf:
            partes.append(munuf)

    if cep_fmt:
        partes.append(f"CEP: {cep_fmt}")

    return ", ".join(partes)


_TIPOS_LOG = {
    "rua", "avenida", "av.", "av", "estrada", "rodovia", "travessa",
    "alameda", "praça", "largo", "vielas", "viela", "rod.", "r.", "r"
}
def _garantir_tipo_LOGRADOURO(prefixo_preferencial: str, endereco: str) -> str:
    e = (endereco or "").strip()
    if not e:
        return e
    primeira = e.split()[0].strip(".,").lower()
    if primeira in _TIPOS_LOG:
        return e
    pref = (prefixo_preferencial or "Rua").strip()
    if not pref.endswith(" "):
        pref += " "
    return pref + e


_UNIDADES = ["", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove"]
_DEZ_A_DEZENOVE = ["dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis", "dezessete", "dezoito", "dezenove"]
_DEZENAS = ["", "", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"]
_CENTENAS = ["", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]

def _centena_por_extenso(n: int) -> str:
    assert 0 <= n <= 999
    if n == 0:
        return ""
    if n == 100:
        return "cem"
    c = n // 100
    d = (n % 100) // 10
    u = n % 10
    partes = []
    if c:
        partes.append(_CENTENAS[c])
    if d == 1:
        if partes: partes.append("e")
        partes.append(_DEZ_A_DEZENOVE[u])
        return " ".join(partes)
    if d >= 2:
        if partes: partes.append("e")
        partes.append(_DEZENAS[d])
    if u:
        if partes: partes.append("e")
        partes.append(_UNIDADES[u])
    return " ".join(partes)

def _grupo_milhar_extenso(n: int, singular: str, plural: str) -> str:
    if n == 0:
        return ""
    if n == 1:
        return f"um {singular}"
    return f"{_centena_por_extenso(n)} {plural}"

def numero_por_extenso_reais(valor: float) -> str:
    if valor < 0:
        return "menos " + numero_por_extenso_reais(-valor)
    inteiro = int(valor)
    centavos = int(round((valor - inteiro) * 100))
    if centavos == 100:
        inteiro += 1
        centavos = 0
    mi = inteiro // 1_000_000
    milhar = (inteiro % 1_000_000) // 1000
    resto = inteiro % 1000

    partes = []
    if mi:
        partes.append(_grupo_milhar_extenso(mi, "milhão", "milhões"))
    if milhar:
        partes.append("mil" if milhar == 1 else f"{_centena_por_extenso(milhar)} mil")
    if resto:
        if partes:
            partes.append("e")
        partes.append(_centena_por_extenso(resto))
    if not partes:
        partes.append("zero")

    reais = "real" if inteiro == 1 else "reais"
    frase = " ".join(partes) + f" {reais}"

    if centavos:
        cent = "centavo" if centavos == 1 else "centavos"
        frase += f" e {_centena_por_extenso(centavos)} {cent}"

    return frase

# =========================================================
#              ARQUIVOS / LOCALIZAÇÃO LOCAL (SEM ONEDRIVE)
# =========================================================
def _app_base_dir() -> str:
    try:
        if getattr(sys, "frozen", False) and hasattr(sys, "executable"):
            return os.path.dirname(sys.executable)
    except Exception:
        pass
    try:
        return os.path.dirname(os.path.abspath(__file__))
    except Exception:
        return os.getcwd()

APP_DIR = _app_base_dir()
DATA_DIR = safe_join(APP_DIR, "data")

def _resolve_local_base() -> str:
    # 1) arquivo local_base.txt ao lado do .exe/.py
    try:
        bases = []
        if getattr(sys, "frozen", False) and hasattr(sys, "executable"):
            bases.append(os.path.dirname(sys.executable))
        try:
            bases.append(os.path.dirname(os.path.abspath(__file__)))
        except Exception:
            pass
        bases.append(os.getcwd())
        for b in bases:
            p = os.path.join(b, "local_base.txt")
            if os.path.exists(p):
                with open(p, "r", encoding="utf-8") as f:
                    base = f.read().strip().strip('"')
                    if base:
                        if not os.path.isabs(base):
                            base = safe_join(APP_DIR, base)
                        os.makedirs(base, exist_ok=True)
                        return base
    except Exception:
        pass

    # 2) variável de ambiente
    base = os.environ.get("LOCAL_FILES_BASE", "").strip()
    if base:
        # Se for relativo, resolve a partir do APP_DIR
        if not os.path.isabs(base):
            base = safe_join(APP_DIR, base)
        try:
            os.makedirs(base, exist_ok=True)
        except Exception:
            pass
        return base
    # 3) padrão
    try:
        os.makedirs(DATA_DIR, exist_ok=True)
    except Exception:
        pass
    return DATA_DIR

LOCAL_BASE = _resolve_local_base()
PASTA_LOGO = safe_join(LOCAL_BASE, "LOGO AUDACES")
PASTA_CONTRATO = safe_join(LOCAL_BASE, "CONTRATO PARA ATUALIZAÇÃO")

# Mantido apenas como fallback local; o app usa DB/API como fonte principal
EXCEL_FILE = safe_join(LOCAL_BASE, "BANCO_DE_DADOS_ORCAMENTO.xlsx")

AUDACES_LOGO_PATH = (
    try_first_existing([
        safe_join(PASTA_LOGO, "audaces.png"),
        safe_join(PASTA_LOGO, "AUDACES.png"),
        r"C:\\Users\\LEANDRO ROSA\\Documents\\Impressões\\LOGO AUDACES\\audaces.png",
        r"C:\\Users\\LEANDRO ROSA\\Documents\\Impressões\\LOGO AUDACES\\AUDACES.png",
    ])
    or find_in_folder(PASTA_LOGO, "audaces", (".png", ".jpg", ".jpeg"))
)

CONTRATO_TEMPLATE = (
    try_first_existing([
        r"C:\\Users\\LEANDRO ROSA\\Documents\\Impressões\\CONTRATO PARA ATUALIZAÇÃO\\CONTRATO COMERCIAL Impressão.docx",
        safe_join(PASTA_CONTRATO, "CONTRATO COMERCIAL Impressão.docx"),
    ])
    or find_in_folder(PASTA_CONTRATO, "CONTRATO COMERCIAL Impressão", (".docx",))
)

ABA_ORCAMENTOS = "Orçamentos"
ABA_CADASTROS = "Cadastros"
ABA_PEDIDOS = "Pedidos"

# =========================================================
#                 PLANILHA EXCEL (CRUD)
# =========================================================
def _header_map(ws):
    headers = [str(c.value).strip() if c.value is not None else "" for c in ws[1]]
    return {h: i for i, h in enumerate(headers)}

def _garantir_headers(ws, headers):
    existing = [cell.value for cell in ws[1]]
    if not any(existing):
        ws.append(headers)
        return
    missing = [h for h in headers if h not in existing]
    if missing:
        new_headers = [x if x else "" for x in existing] + missing
        for c in range(1, len(new_headers) + 1):
            ws.cell(row=1, column=c, value=new_headers[c - 1])
        target_len = len(new_headers)
        for r in range(2, ws.max_row + 1):
            row_len = len([c.value for c in ws[r]])
            for _ in range(target_len - row_len):
                ws.cell(row=r, column=row_len + 1, value=None)

def init_excel():
    os.makedirs(os.path.dirname(EXCEL_FILE), exist_ok=True)

    headers_orc = [
        "ID Orçamento","Data/Hora","Tipo de Serviço",
        "CLIENTE (Etiqueta PDF)","CLIENTE (Valor)","Documento","CNPJ/CPF","E-mail",
        "Vendedor",  # NOVO
        "Status","Quantidade","Unidade","Metros","Preço por metro",
        "Forma de Pagamento",           # NOVO no Orçamento
        "Valor Total",
    ]
    headers_cad = [
        "Documento","CNPJ/CPF","Razão Social/Nome","Nome Fantasia","Contato",
        "Inscrição Estadual","Situação IE","Inscrição Municipal","Situação IM","Situação Cadastral",
        "CEP","Endereço","Número","Complemento","Bairro","Município","UF",
        "Telefone 1","Telefone 2","E-mail (CNPJ)","E-mail (Manual)",
        "Entrega CEP","Entrega Endereço","Entrega Número","Entrega Complemento","Entrega Bairro",
        "Entrega Município","Entrega UF",
        "Desconto Duração","Desconto Unidade",
        "Criado em","Atualizado em",
    ]
    headers_pedidos = [
        "ID","Pedido","Tipo de Serviço","Status do CLIENTE","Quantidade (m)","Valor Unitário","Valor Total",
        "Data/Hora da criação do pedido","ID Orçamento","Documento","CNPJ/CPF","CLIENTE",
        "Vendedor",  # NOVO
        "Forma de Pagamento Orçamento",  # NOVO
        "Forma de Pagamento Contrato",   # NOVO
        "% Comissão Vendedor",           # NOVO
        "Valor Comissão Vendedor",       # NOVO
        "% Comissão ADM",                # NOVO
        "Valor Comissão ADM",            # NOVO
    ]

    if not os.path.exists(EXCEL_FILE):
        wb = Workbook()
        ws = wb.active
        ws.title = ABA_ORCAMENTOS
        ws.append(headers_orc)
        wb.create_sheet(ABA_CADASTROS).append(headers_cad)
        wb.create_sheet(ABA_PEDIDOS).append(headers_pedidos)
        wb.save(EXCEL_FILE)
    else:
        try:
            wb = load_workbook(EXCEL_FILE)
        except zipfile.BadZipFile:
            # Arquivo existente, mas inválido (ex.: placeholder do OneDrive ou arquivo corrompido)
            try:
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                backup = f"{EXCEL_FILE}.backup_{ts}"
                try:
                    shutil.copy2(EXCEL_FILE, backup)
                except Exception:
                    pass
                wb = Workbook()
                ws = wb.active
                ws.title = ABA_ORCAMENTOS
                ws.append(headers_orc)
                wb.create_sheet(ABA_CADASTROS).append(headers_cad)
                wb.create_sheet(ABA_PEDIDOS).append(headers_pedidos)
                wb.save(EXCEL_FILE)
            except Exception:
                raise
        ws = wb[ABA_ORCAMENTOS] if ABA_ORCAMENTOS in wb.sheetnames else wb.create_sheet(ABA_ORCAMENTOS, 0)
        _garantir_headers(ws, headers_orc)
        ws2 = wb[ABA_CADASTROS] if ABA_CADASTROS in wb.sheetnames else wb.create_sheet(ABA_CADASTROS)
        _garantir_headers(ws2, headers_cad)
        ws3 = wb[ABA_PEDIDOS] if ABA_PEDIDOS in wb.sheetnames else wb.create_sheet(ABA_PEDIDOS)
        _garantir_headers(ws3, headers_pedidos)
        if "Vendas" in wb.sheetnames:
            del wb["Vendas"]
        wb.save(EXCEL_FILE)


def load_wb_safe(path: str, **kwargs):
    try:
        return load_workbook(path, **kwargs)
    except zipfile.BadZipFile:
        # Tenta recuperar recriando/normalizando a planilha
        init_excel()
        return load_workbook(path, **kwargs)
def salvar_excel_orcamento(dados):
    # Agora salva via API/DB. Mantemos a assinatura para mínimo impacto.
    try:
        body = {
            "tipo_servico": str(dados.get("Tipo de Serviço") or ""),
            "CLIENTE": str(dados.get("CLIENTE (Valor)") or ""),
            "cnpj": str(dados.get("CNPJ/CPF") or ""),
            "email": str(dados.get("E-mail") or ""),
            "status": str(dados.get("Status") or ""),
            "unidade": str(dados.get("Unidade") or ""),
            "quantidade": str(dados.get("Quantidade") or ""),
        }
        return api_post("/api/orcamentos", body)
    except Exception as ex:
        # Como fallback (somente para desenvolvimento), ainda tenta local
        wb = load_wb_safe(EXCEL_FILE)
        ws = wb[ABA_ORCAMENTOS]
        if isinstance(dados, dict):
            hmap = _header_map(ws)
            row = [dados.get(h, "") for h in hmap.keys()]
            ws.append(row)
        else:
            ws.append(dados)
        wb.save(EXCEL_FILE)
        return {"ok": True, "fallback": str(ex)}

def salvar_excel_cadastro(dados_dict):
    try:
        return api_post("/api/cadastros", dict(dados_dict))
    except Exception as ex:
        wb = load_wb_safe(EXCEL_FILE)
        ws = wb[ABA_CADASTROS]
        hmap = _header_map(ws)
        agora = data_hora_tokens()["combinado"]
        dados = dict(dados_dict)
        dados.setdefault("Criado em", agora)
        dados["Atualizado em"] = agora
        row = [dados.get(h, "") for h in hmap.keys()]
        ws.append(row)
        wb.save(EXCEL_FILE)
        return {"ok": True, "fallback": str(ex)}

def atualizar_excel_cadastro(doc_formatado: str, dados_dict: dict) -> bool:
    wb = load_wb_safe(EXCEL_FILE)
    ws = wb[ABA_CADASTROS]
    hmap = _header_map(ws)
    if "CNPJ/CPF" not in hmap:
        wb.close()
        return False
    col_doc = hmap["CNPJ/CPF"] + 1
    alvo_row, alvo_digits = None, re.sub(r"\D", "", doc_formatado or "")
    for r in range(2, ws.max_row + 1):
        val = ws.cell(row=r, column=col_doc).value
        if val and re.sub(r"\D", "", str(val)) == alvo_digits:
            alvo_row = r
            break
    if not alvo_row:
        wb.close()
        return False
    criado_col = hmap.get("Criado em")
    criado_atual = ws.cell(row=alvo_row, column=criado_col + 1).value if criado_col is not None else ""
    agora = data_hora_tokens()["combinado"]
    for nome, idx0 in hmap.items():
        if nome == "Criado em":
            ws.cell(row=alvo_row, column=idx0 + 1, value=criado_atual or dados_dict.get("Criado em", ""))
        elif nome == "Atualizado em":
            ws.cell(row=alvo_row, column=idx0 + 1, value=agora)
        else:
            ws.cell(row=alvo_row, column=idx0 + 1, value=dados_dict.get(nome, ws.cell(row=alvo_row, column=idx0 + 1).value))
    wb.save(EXCEL_FILE)
    return True

def buscar_cadastro_por_documento(tipo: str, valor_digitado: str) -> dict | None:
    # Preferir API/DB
    try:
        digits = re.sub(r"\D", "", valor_digitado or "")
        resp = api_get(f"/api/cadastros?cnpj={digits}")
        rows = resp.get("rows") or []
        if rows:
            return rows[0]
    except Exception:
        pass
    # Fallback Excel local, se existir
    if not os.path.exists(EXCEL_FILE):
        return None
    try:
        wb = load_wb_safe(EXCEL_FILE, read_only=True, data_only=True)
        if ABA_CADASTROS not in wb.sheetnames:
            wb.close()
            return None
        ws = wb[ABA_CADASTROS]
        hmap = _header_map(ws)
        if "CNPJ/CPF" not in hmap:
            wb.close()
            return None
        alvo = re.sub(r"\D", "", formatar_doc(tipo, valor_digitado))
        col = hmap["CNPJ/CPF"]
        ultimo = None
        for row in ws.iter_rows(min_row=2, values_only=True):
            v = row[col] if col < len(row) else None
            if v and re.sub(r"\D", "", str(v)) == alvo:
                d = {n: (row[i] if i < len(row) else None) for n, i in hmap.items()}
                ultimo = d
        wb.close()
        return ultimo
    except Exception:
        return None

def get_orcamento_by_id(id_orc: str) -> dict | None:
    # API primeiro
    try:
        resp = api_get(f"/api/orcamentos?id={urllib.parse.quote(id_orc or '')}")
        rows = resp.get("rows") or []
        if rows:
            return rows[0]
    except Exception:
        pass
    # Fallback Excel
    if not os.path.exists(EXCEL_FILE):
        return None
    wb = load_wb_safe(EXCEL_FILE, read_only=True, data_only=True)
    if ABA_ORCAMENTOS not in wb.sheetnames:
        wb.close()
        return None
    ws = wb[ABA_ORCAMENTOS]
    hmap = _header_map(ws)
    for row in ws.iter_rows(min_row=2, values_only=True):
        d = {n: (row[i] if i < len(row) else None) for n, i in hmap.items()}
        if str(d.get("ID Orçamento") or "").strip() == (id_orc or "").strip():
            wb.close()
            return d
    wb.close()
    return None

def get_orcamentos_list(doc_formatado: str | None = None, id_orc: str | None = None) -> list[dict]:
    # Primeiro tenta pela API (já no formato de labels)
    try:
        params = []
        if id_orc:
            params.append(f"id={urllib.parse.quote(id_orc)}")
        if doc_formatado:
            params.append(f"cnpj={re.sub(r'\\D','', doc_formatado)}")
        q = ("?" + "&".join(params)) if params else ""
        resp = api_get(f"/api/orcamentos{q}")
        rows = resp.get("rows") or []
        if rows:
            return rows
    except Exception:
        pass
    out = []
    def _looks_currency_ptbr(s: str) -> bool:
        t = (s or "").strip()
        return bool(re.match(r"^\d{1,3}(?:\.\d{3})*,\d{2}$", t))
    def _looks_number_ptbr(s: str) -> bool:
        t = (s or "").strip()
        return bool(re.match(r"^\d{1,3}(?:\.\d{3})*,\d{1,3}$", t))
    def _normalize_row(d: dict) -> dict:
        outd = dict(d)
        # Nome do CLIENTE pode ter sido corrompido por gravação antiga; tenta recuperar pelo cadastro
        nome_cli = extrair_nome_CLIENTE(outd)
        if not nome_cli or _looks_currency_ptbr(str(nome_cli)):
            doc_tipo = str(outd.get("Documento") or "")
            doc_val = str(outd.get("CNPJ/CPF") or "")
            cad = buscar_cadastro_por_documento(doc_tipo, doc_val) or {}
            nome_corrigido = extrair_nome_CLIENTE(cad)
            if nome_corrigido:
                outd["CLIENTE (Valor)"] = nome_corrigido

        # Normaliza Preço por metro
        preco_txt = str(outd.get("Preço por metro") or "")
        preco_num = _parse_ptbr_float(preco_txt)

        # Metros: se vier texto (ex.: "Centímetros") ou inválido, tenta recompor
        metros_txt = str(outd.get("Metros") or "")
        metros_num = _parse_ptbr_float(metros_txt)
        if not _looks_number_ptbr(metros_txt):
            # a) pela Quantidade + Unidade
            qtd_txt = str(outd.get("Quantidade") or "")
            uni = str(outd.get("Unidade") or "")
            qtd_num = _parse_ptbr_float(qtd_txt)
            if qtd_num > 0:
                metros_num = (qtd_num / 100.0) if uni.lower().startswith("cent") else qtd_num
            # b) pela divisão Valor Total / Preço por metro
            if metros_num <= 0 and preco_num > 0:
                vt_num = _parse_ptbr_float(str(outd.get("Valor Total") or "0"))
                if vt_num > 0:
                    metros_num = vt_num / preco_num
            if metros_num > 0:
                outd["Metros"] = format_num_ptbr(metros_num)

        # Valor Total: se vazio/ inválido, calcula pela multiplicação
        vtotal_txt = str(outd.get("Valor Total") or "")
        if not _looks_currency_ptbr(vtotal_txt):
            if metros_num <= 0:
                metros_num = _parse_ptbr_float(str(outd.get("Metros") or "0"))
            if metros_num > 0 and preco_num > 0:
                outd["Valor Total"] = format_num_ptbr(metros_num * preco_num)

        return outd
    if not os.path.exists(EXCEL_FILE):
        return out
    wb = load_wb_safe(EXCEL_FILE, read_only=True, data_only=True)
    if ABA_ORCAMENTOS not in wb.sheetnames:
        wb.close()
        return out
    ws = wb[ABA_ORCAMENTOS]
    hmap = _header_map(ws)
    for row in ws.iter_rows(min_row=2, values_only=True):
        d = {n: (row[i] if i < len(row) else None) for n, i in hmap.items()}
        if id_orc and str(d.get("ID Orçamento") or "").strip() != id_orc.strip():
            continue
        if doc_formatado and str(d.get("CNPJ/CPF") or "").strip() != doc_formatado.strip():
            continue
        out.append(_normalize_row(d))
    wb.close()
    return out

def _parse_datetime_ptbr(txt: str) -> datetime | None:
    try:
        return datetime.strptime((txt or "").strip(), "%d/%m/%Y %H:%M:%S")
    except Exception:
        return None

def get_ultimo_pedido_data(doc_formatado: str) -> datetime | None:
    try:
        digits = re.sub(r"\D", "", doc_formatado or "")
        resp = api_get(f"/api/pedidos?cnpj={digits}")
        rows = resp.get("rows") or []
        def parse_row(r):
            txt = r.get("Data/Hora da criação do pedido") or r.get("data_hora_criacao") or ""
            try:
                return datetime.strptime(str(txt), "%d/%m/%Y %H:%M:%S")
            except Exception:
                return None
        dts = [d for d in (parse_row(r) for r in rows) if d]
        return max(dts) if dts else None
    except Exception:
        pass
    if not os.path.exists(EXCEL_FILE):
        return None
    try:
        wb = load_wb_safe(EXCEL_FILE, read_only=True, data_only=True)
        if ABA_PEDIDOS not in wb.sheetnames:
            wb.close()
            return None
        ws = wb[ABA_PEDIDOS]
        hmap = _header_map(ws)
        idx_doc = hmap.get("CNPJ/CPF")
        idx_dh = hmap.get("Data/Hora da criação do pedido")
        if idx_doc is None or idx_dh is None:
            wb.close()
            return None
        alvo = re.sub(r"\D", "", doc_formatado or "")
        ultimo = None
        for row in ws.iter_rows(min_row=2, values_only=True):
            vdoc = row[idx_doc] if idx_doc < len(row) else None
            vdh = row[idx_dh] if idx_dh < len(row) else None
            if vdoc and re.sub(r"\D", "", str(vdoc)) == alvo:
                d = _parse_datetime_ptbr(str(vdh)) if isinstance(vdh, str) else (vdh if isinstance(vdh, datetime) else None)
                if d and (ultimo is None or d > ultimo):
                    ultimo = d
        wb.close()
        return ultimo
    except Exception:
        return None

def desconto_automatico_por_pedido(doc_formatado: str) -> bool:
    cad = buscar_cadastro_por_documento("CNPJ" if len(re.sub(r"\D","", doc_formatado))==14 else "CPF", doc_formatado) or {}
    dur_txt = str(cad.get("Desconto Duração") or "").strip()
    unid = str(cad.get("Desconto Unidade") or "").strip().lower()
    try:
        dur = int(dur_txt)
    except Exception:
        return False
    if dur <= 0 or unid not in ("meses","anos"):
        return False
    ref = get_ultimo_pedido_data(doc_formatado)
    if not ref:
        return False
    dias = dur * (30 if unid == "meses" else 365)
    limite = ref + timedelta(days=dias)
    return datetime.now() <= limite

def get_proximo_sequencial(sigla: str) -> int:
    try:
        tipo = "Impressão" if sigla == "IM" else "Digitalização"
        resp = api_get(f"/api/proximo-id?tipo_servico={urllib.parse.quote(tipo)}")
        _ = resp.get("id")
        return 1
    except Exception:
        pass
    if not os.path.exists(EXCEL_FILE):
        return 1
    wb = load_wb_safe(EXCEL_FILE, read_only=True, data_only=True)
    if ABA_ORCAMENTOS not in wb.sheetnames:
        wb.close()
        return 1
    ws = wb[ABA_ORCAMENTOS]
    cont = 0
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] and str(row[0]).startswith(f"OR-{sigla}"):
            cont += 1
    wb.close()
    return cont + 1

def gerar_id(tipo_servico: str) -> str:
    if not tipo_servico:
        return ""
    try:
        resp = api_get(f"/api/proximo-id?tipo_servico={urllib.parse.quote(tipo_servico)}")
        return str(resp.get("id") or "")
    except Exception:
        sigla = "IM" if tipo_servico.startswith("Imp") else "DG"
        return f"OR-{sigla}{get_proximo_sequencial(sigla)}{data_hora_tokens()['data_compacta']}"

def get_proximo_pedido_numero() -> int:
    """Pega o MAIOR valor de 'Pedido' via API; fallback Excel."""
    try:
        resp = api_get("/api/pedidos")
        rows = resp.get("rows") or []
        maior = 0
        for r in rows:
            try:
                n = int(str(r.get("Pedido") or r.get("pedido") or "0").strip())
                if n > maior:
                    maior = n
            except Exception:
                pass
        return maior + 1
    except Exception:
        pass
    if not os.path.exists(EXCEL_FILE):
        return 1
    wb = load_workbook(EXCEL_FILE, read_only=True, data_only=True)
    if ABA_PEDIDOS not in wb.sheetnames:
        wb.close()
        return 1
    ws = wb[ABA_PEDIDOS]
    hmap = _header_map(ws)
    idx = hmap.get("Pedido", 1)
    maior = 0
    for row in ws.iter_rows(min_row=2, values_only=True):
        val = row[idx] if idx < len(row) else None
        try:
            n = int(str(val).strip())
            if n > maior:
                maior = n
        except Exception:
            pass
    wb.close()
    return maior + 1

def salvar_excel_pedido(dados_dict):
    try:
        return api_post("/api/pedidos", dict(dados_dict))
    except Exception as ex:
        wb = load_wb_safe(EXCEL_FILE)
        ws = wb[ABA_PEDIDOS]
        hmap = _header_map(ws)
        row = [dados_dict.get(h, "") for h in hmap.keys()]
        ws.append(row)
        wb.save(EXCEL_FILE)
        return {"ok": True, "fallback": str(ex)}

# =========================================================
#                         PDF Orçamento
# =========================================================
def gerar_pdf_orcamento(dados, pasta_destino):
    (
        id_orc, datahora, tipo_servico, CLIENTE_label, CLIENTE_valor,
        documento, doc_valor, email, vendedor, status, qtd, unidade, metros, preco,
        forma_pgto,  # NOVO
        total,
    ) = dados

    nome_arquivo = f"{sanitize_filename(id_orc)}_{sanitize_filename(CLIENTE_valor)}_{datahora.replace('/', '-').replace(':', '-')}.pdf"
    caminho = os.path.join(pasta_destino, nome_arquivo)
    c = canvas.Canvas(caminho, pagesize=A4)
    W, H = A4

    c.setFont(FONT_REG, 10)
    c.setFillColor(colors.black)
    try:
        cab = datetime.strptime(datahora, "%d/%m/%Y %H:%M:%S").strftime("%d/%m/%Y, %H:%M")
    except Exception:
        cab = datahora
    y_top = H - (MARGEM * 0.35)
    c.drawString(MARGEM + HEADER_LEFT_OFFSET, y_top, cab)
    c.drawCentredString(W / 2 + HEADER_CENTER_OFFSET, y_top, "Fashion Tech - Audaces RJ e ES")

    titulo_y_base = H - (MARGEM + 10)
    try:
        if AUDACES_LOGO_PATH and os.path.exists(AUDACES_LOGO_PATH):
            logo_x = W - LOGO_RIGHT_MARGIN - LOGO_WIDTH
            logo_y = H - LOGO_TOP_MARGIN - LOGO_HEIGHT
            c.drawImage(
                AUDACES_LOGO_PATH, logo_x, logo_y,
                width=LOGO_WIDTH, height=LOGO_HEIGHT,
                preserveAspectRatio=True, mask="auto",
            )
            titulo_y_base = min(titulo_y_base, logo_y - 18)
    except Exception:
        pass

    c.setFont(FONT_BOLD, H1_SIZE)
    c.setFillColor(AZUL)
    # Se nao houve logo, garante espaco adequado acima
    y = titulo_y_base
    c.drawCentredString(W / 2, y, "Fashion Tech - Audaces RJ e ES")
    y -= H1_SIZE + 2
    c.drawCentredString(W / 2, y, "Orçamento de Impressão de Riscos")

    y -= TITLE_TO_BODY_GAP
    x = MARGEM
    c.setFillColor(colors.black)
    c.setFont(FONT_REG, BODY_SIZE)

    def draw_label_value(label: str, value: str):
        nonlocal y
        c.setFont(FONT_BOLD, BODY_SIZE)
        c.drawString(x, y, label + ":")
        lw = c.stringWidth(label + ":", FONT_BOLD, BODY_SIZE)
        c.setFont(FONT_REG, BODY_SIZE)
        c.drawString(x + lw + 6, y, value)
        y -= LINE

    qtd_num = _parse_ptbr_float(qtd)
    unidade_display = "Metros" if (unidade == "Metro" and abs(qtd_num - 1.0) > 1e-9) else unidade
    metros_num = _parse_ptbr_float(metros)
    metros_unit = "metros" if abs(metros_num - 1.0) > 1e-9 else "metro"

    # Exibir Vendedor antes do ID, mantendo o espaçamento inicial do bloco
    draw_label_value("Vendedor", vendedor or "-")
    draw_label_value("ID Orçamento", id_orc)
    draw_label_value("Tipo de Serviço", tipo_servico)
    draw_label_value(CLIENTE_label, CLIENTE_valor)
    draw_label_value("CPF" if documento == "CPF" else "CNPJ", doc_valor)
    draw_label_value("E-mail", email)
    draw_label_value("Qtde.", f"{qtd} {unidade_display}")
    if unidade and unidade.lower().startswith("cent"):
        draw_label_value("Convertido", f"{metros} {metros_unit}")
    draw_label_value("Preço por metro", f"R$ {preco}")
    draw_label_value("Forma de Pagamento", forma_pgto if (forma_pgto or "").strip() else "-")

    y -= LINE
    c.setFont(FONT_BOLD, TOTAL_SIZE)
    c.setFillColor(AZUL)
    c.drawString(x, y, f"Valor Total: R$ {total}")
    c.setFillColor(colors.black)
    y -= 2 * LINE

    c.setFont(FONT_ITAL, BODY_SIZE)
    c.drawString(x, y, f"Gerado em: {datahora}")
    y -= 2 * LINE

    c.setFont(FONT_REG, BODY_SIZE)
    label = "Assinatura do CLIENTE:"
    c.drawString(x, y, label)
    lw = c.stringWidth(label, FONT_REG, BODY_SIZE)
    line_y = y - 3
    c.setLineWidth(1)
    c.line(x + lw + 12, line_y, W - MARGEM, line_y)

    y -= 2 * LINE
    c.setFont(FONT_REG, 11)
    c.drawString(x, y, "Fashion Tech - Audaces RJ e ES")
    y -= LINE * 0.9
    c.drawString(x, y, "E-mail: fashiontech.impressao@audaces.com | Telefone: (21) 99132-3562")
    y -= LINE * 0.9
    c.drawString(x, y, "Leandro Rosa / Supervisor ADM")

    y -= 2 * LINE
    c.drawString(x, y, "Especificações Técnicas - Plotter Audaces Essence 185")
    y -= LINE * 1.1
    c.setFont(FONT_REG, BODY_SIZE)
    bullets = [
        "Largura máxima de plotagem: 185 cm",
        "área útil de Impressão: 170 cm",
        "Tecnologia de Impressão: Inkjet (jato de tinta)",
        "Principais aplicações: Impressão de encaixes para corte de tecido em confecções",
    ]
    for b in bullets:
        c.drawString(x + 16, y, f"- {b}")
        y -= LINE * 0.9

    c.save()
    return caminho

# =========================================================
#                 CONTRATO .DOCX  (python-docx)
# =========================================================
def _docx_search_replace_preservando_formatacao(document, mapping: dict):
    """
    Substitui placeholders preservando formatação (bold/italic/underline/fonte/tamanho) dos runs afetados.
    Estratégia:
      - Para cada parágrafo/célula, varre a sequência de runs concatenando o texto.
      - Quando encontra um placeholder (que pode atravessar runs), substitui por um novo run:
           * texto = replacement
           * estilo = copia do primeiro run do match (inclui bold/italic/underline/fonte/size)
      - Mantém os runs restantes.
    """
    from docx.text.paragraph import Paragraph
    from docx.table import _Cell
    from docx.oxml.text.run import CT_R
    from docx.text.run import Run

    def _copy_run_style(src: Run, dst: Run):
        try:
            dst.bold = src.bold
            dst.italic = src.italic
            dst.underline = src.underline
            if src.style:
                dst.style = src.style
            if src.font:
                dst.font.name = src.font.name
                dst.font.size = src.font.size
                dst.font.color.rgb = getattr(src.font.color, "rgb", None)
        except Exception:
            pass

    def _replace_in_runs(par: Paragraph):
        if not mapping:
            return
        if not par.runs:
            return

        # Concatena textos e mapeia índices -> (run_index, offset_no_run)
        texto_total = "".join([r.text or "" for r in par.runs])
        if not texto_total:
            return

        for alvo, repl in mapping.items():
            if not alvo:
                continue
            repl_str = "" if repl is None else str(repl)
            start = 0
            while True:
                idx = texto_total.find(alvo, start)
                if idx == -1:
                    break

                # Calcula quais runs cobrem [idx, idx+len(alvo))
                alvo_end = idx + len(alvo)
                acumulado = 0
                primeiro_i = None
                ultimo_i = None
                off_in_first = 0
                off_in_last_end = 0

                for i, r in enumerate(par.runs):
                    txt = r.text or ""
                    ini = acumulado
                    fim = acumulado + len(txt)
                    if primeiro_i is None and idx < fim:
                        primeiro_i = i
                        off_in_first = idx - ini
                    if primeiro_i is not None and alvo_end <= fim:
                        ultimo_i = i
                        off_in_last_end = alvo_end - ini
                        break
                    acumulado = fim

                if primeiro_i is None or ultimo_i is None:
                    # Falhou o mapeamento; evita loop infinito
                    start = alvo_end
                    continue

                # Partes antes e depois do match dentro dos runs extremos
                try:
                    run_first = par.runs[primeiro_i]
                    run_last = par.runs[ultimo_i]
                except Exception:
                    # Fallback seguro: substitui no parágrafo inteiro sem preservar formatação
                    try:
                        par.text = (par.text or "").replace(alvo, repl_str)
                    except Exception:
                        pass
                    texto_total = texto_total[:idx] + repl_str + texto_total[alvo_end:]
                    start = idx + len(repl_str)
                    continue

                prefixo = (run_first.text or "")[:off_in_first]
                sufixo = (run_last.text or "")[off_in_last_end:]

                # Apaga runs do intervalo e recria com (prefixo) + [repl como novo run] + (sufixo)
                # 1) Ajusta texto do primeiro run para o prefixo
                run_first.text = prefixo

                # 2) Remove runs intermediários COMPLETOS entre first e last
                for _ in range(max(0, ultimo_i - primeiro_i - 1)):
                    if len(par.runs) <= primeiro_i + 1:
                        break
                    rmid = par.runs[primeiro_i + 1]
                    rmid._element.getparent().remove(rmid._element)

                # 3) Ajusta texto do último run para sufixo
                if ultimo_i != primeiro_i:
                    run_last = par.runs[primeiro_i + 1]  # após remoções, last virou o próximo
                    run_last.text = sufixo
                else:
                    # quando match inteiro estava dentro de um único run
                    run_last.text = sufixo

                # 4) Simplifica: aplica prefixo + replacement + sufixo no primeiro run
                #    e remove todos os runs subsequentes até o 'last'.
                run_first.text = (prefixo or "") + repl_str + (sufixo or "")
                for _ in range(max(0, ultimo_i - primeiro_i)):
                    if len(par.runs) <= primeiro_i + 1:
                        break
                    rdel = par.runs[primeiro_i + 1]
                    rdel._element.getparent().remove(rdel._element)

                # Recalcula texto_total para próximos matches (ATUALIZAÇÃO local)
                texto_total = texto_total[:idx] + repl_str + texto_total[alvo_end:]
                start = idx + len(repl_str)

    def _walk_block(block):
        if isinstance(block, Paragraph):
            _replace_in_runs(block)
        else:
            # Tabelas
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_runs(p)
                    for t2 in cell.tables:
                        _walk_block(t2)

    # parágrafos soltos
    for p in document.paragraphs:
        _replace_in_runs(p)
    # Tabelas
    for t in document.tables:
        _walk_block(t)

def gerar_contrato_docx(contexto: dict, pasta_destino: str) -> tuple[str, str]:
    try:
        from docx import Document  # python-docx
    except Exception as ex:
        return "", f"Módulo python-docx não disponível: {ex}"

    if not CONTRATO_TEMPLATE or not os.path.exists(CONTRATO_TEMPLATE):
        return "", "Template do contrato não encontrado no caminho configurado. Verifique LOCAL_FILES_BASE ou o caminho fixo em Documentos/Impressões."

    try:
        doc = Document(CONTRATO_TEMPLATE)
    except Exception as ex:
        return "", ("Falha ao abrir o template. Detalhe: " + str(ex))

    valor_unit_docx = "R$ " + (contexto.get("valor_unit", "") or "").strip()
    valor_total_docx = "R$ " + (contexto.get("valor_total", "") or "").strip()
    data_hoje = datetime.now().strftime("%d/%m/%Y")
    forma_pgto = (contexto.get("forma_pgto", "") or "").strip()
    try:
        valor_total_float = _parse_ptbr_float(contexto.get("valor_total", "0") or "0")
    except Exception:
        valor_total_float = 0.0
    valor_total_extenso = numero_por_extenso_reais(valor_total_float)
    end_concat_original = contexto.get("empresa_endereco_concat", "") or ""
    end_concat_com_tipo = _garantir_tipo_LOGRADOURO("Rua", end_concat_original)

    mapping = {
        "(INCLUIR Razão SOCIAL DO CLIENTE)": contexto["CLIENTE"],
        "(INCLUIR CNPJ OU CPF DO CLIENTE)": contexto["doc_valor"],
        "(INCLUIR Endereço DE ENTREGA DO CLIENTE)": contexto["end_entrega"],
        "+55 (INCLUIR Número DE TELEFONE DO CLIENTE)": f"+55 {contexto['telefone']}",
        "(INCLUIR E-mail do CLIENTE)": contexto["email"],
        "INCLUIR Razão SOCIAL DA EMPRESA": contexto["empresa_razao"],
        "INCLUIR Número DO CNPJ DA EMPRESA": contexto["empresa_cnpj"],
        "INCLUIR Endereço DA EMPRESA COMPLETO CONCATENADO": end_concat_com_tipo,
        "EDITAR DATA": data_hoje,
        "(EDITAR DATA)": data_hoje,
        "(INCLUIR FORMA DE PAGAMENTO)": forma_pgto,
        "(FORMA DE PAGAMENTO)": forma_pgto,
        "FORMA DE PAGAMENTO": forma_pgto,
        "(INCLUIR VALOR ESCRITO POR EXTENSO)": valor_total_extenso,
        "TIPO Serviço": contexto.get("tipo_servico", ""),
        "TOTAL EM METROS": contexto.get("total_metros", ""),
        "VALOR UNIT.": valor_unit_docx,
        "VALOR UNIT": valor_unit_docx,
        "VALOR TOTAL": valor_total_docx,
    }

    # Normaliza valores do mapeamento para string (evita None em Run.text)
    mapping = {k: ("" if v is None else str(v)) for k, v in mapping.items()}

    # >>> Preserva negrito/estilos <<<
    try:
        _docx_search_replace_preservando_formatacao(doc, mapping)
    except Exception as ex:
        return "", f"Falha ao processar o template do contrato: {ex}"

    ts = datetime.now().strftime("%d-%m-%Y %H-%M-%S")
    id_orc = contexto["id_orc"]
    pos = id_orc.find("-")
    sufixo = id_orc[pos + 1 :] if pos >= 0 else id_orc
    CLIENTE_nome = sanitize_filename(contexto.get('CLIENTE', ''))
    sufixo_sanit = sanitize_filename(sufixo)
    nome_arquivo = f"Contrato Impressão- {sufixo_sanit}_{CLIENTE_nome}_{ts}.docx"
    caminho = os.path.join(pasta_destino, nome_arquivo)
    try:
        doc.save(caminho)
    except Exception as ex:
        return "", f"não foi possível salvar o DOCX: {ex}"
    return caminho, ""

def converter_contrato_para_pdf(caminho_docx: str, pasta_destino: str) -> tuple[str, str]:
    out_pdf = os.path.splitext(os.path.join(pasta_destino, os.path.basename(caminho_docx)))[0] + ".pdf"
    try:
        from docx2pdf import convert  # type: ignore  # pip install docx2pdf
        convert(caminho_docx, out_pdf)
        return out_pdf, ""
    except Exception as ex:
        # Fallback via automação do Word (pywin32)
        try:
            import pythoncom
            pythoncom.CoInitialize()
            import win32com.client  # type: ignore
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(caminho_docx)
            wdExportFormatPDF = 17
            # 0=wdExportOptimizeForPrint, 1=wdExportOptimizeForOnScreen
            doc.ExportAsFixedFormat(out_pdf, wdExportFormatPDF, OpenAfterExport=False)
            doc.Close(False)
            word.Quit()
            return out_pdf, ""
        except Exception as ex2:
            try:
                # Garante fechamento de instância do Word mesmo em falhas
                word.Quit()
            except Exception:
                pass
            msg = (
                "Erro ao gerar PDF. Necessário docx2pdf ou Microsoft Word (COM) disponível. "
                f"Detalhe1: {ex} | Detalhe2: {ex2}"
            )
            return "", msg

# =========================================================
#                           APP
# =========================================================
def main(page: ft.Page):
    page.title = "Fashion Tech - Audaces RJ e ES - Orçamento de Impressão de Riscos"
    page.scroll = "adaptive"
    pill = ft.ButtonStyle(shape=ft.RoundedRectangleBorder(radius=20))

    # NAV
    resultado_global = ft.Text("", size=14, weight="bold")

    # ===== Login simples =====
    current_user = {"usuario": None, "nome": None, "is_admin": False, "permissoes": "*"}
    login_user = ft.TextField(label="Usuário", width=200)
    login_pass = ft.TextField(label="Senha", password=True, can_reveal_password=True, width=200)
    login_msg = ft.Text("")

    def _do_login(e=None):
        try:
            r = api_post("/api/login", {"usuario": login_user.value or "", "senha": login_pass.value or ""})
            current_user.update(r)
            page.dialog.open = False
            # após login, mostra Módulo Orçamentos como padrão
            try:
                show_orc()
            except Exception:
                pass
            page.update()
        except Exception as ex:
            login_msg.value = f"Falha no login: {ex}"
            page.update()

    def require_admin(on_ok):
        usr = ft.TextField(label="Usuário (ADM)", width=200)
        pwd = ft.TextField(label="Senha", password=True, can_reveal_password=True, width=200)
        msg = ft.Text("")
        def _check(e=None):
            try:
                r = api_post("/api/login", {"usuario": usr.value or "", "senha": pwd.value or ""})
                if not r.get("is_admin"):
                    msg.value = "Permissão negada (não  administrador)."; page.update(); return
                page.dialog.open = False; page.update(); on_ok()
            except Exception as ex:
                msg.value = f"Falha: {ex}"; page.update()
        page.dialog = ft.AlertDialog(modal=True, title=ft.Text("Confirmação de administrador"), content=ft.Column([usr, pwd, msg], tight=True), actions=[ft.ElevatedButton("OK", on_click=_check, style=pill)], actions_alignment=ft.MainAxisAlignment.END)
        page.dialog.open = True; page.update()

    page.dialog = ft.AlertDialog(
        modal=True,
        title=ft.Text("Login"),
        content=ft.Column([login_user, login_pass, login_msg], tight=True),
        actions=[ft.ElevatedButton("Entrar", on_click=_do_login, style=pill)],
        actions_alignment=ft.MainAxisAlignment.END,
    )
    page.dialog.open = True
    page.update()

    def show_orc(e=None):
        cont_orcamento.visible = True
        cont_cadastro.visible = False
        cont_contrato.visible = False
        try:
            cont_relatorios.visible = False
            cont_usuarios.visible = False
        except Exception:
            pass
        resultado_global.value = ""
        page.update()

    def show_cad(e=None):
        cont_orcamento.visible = False
        cont_cadastro.visible = True
        cont_contrato.visible = False
        try:
            cont_relatorios.visible = False
            cont_usuarios.visible = False
        except Exception:
            pass
        resultado_global.value = ""
        page.update()

    def show_contrato(e=None):
        cont_orcamento.visible = False
        cont_cadastro.visible = False
        cont_contrato.visible = True
        try:
            cont_relatorios.visible = False
            cont_usuarios.visible = False
        except Exception:
            pass
        resultado_global.value = ""
        page.update()

    # Placeholders para containers que serão definidos mais abaixo
    def show_relatorios(e=None):
        cont_orcamento.visible = False
        cont_cadastro.visible = False
        cont_contrato.visible = False
        if 'cont_relatorios' in locals():
            cont_relatorios.visible = True
        if 'cont_usuarios' in locals():
            cont_usuarios.visible = False
        resultado_global.value = ""
        page.update()

    def show_usuarios(e=None):
        cont_orcamento.visible = False
        cont_cadastro.visible = False
        cont_contrato.visible = False
        if 'cont_relatorios' in locals():
            cont_relatorios.visible = False
        if 'cont_usuarios' in locals():
            cont_usuarios.visible = True
        resultado_global.value = ""
        page.update()

    nav = ft.Row(
        [
            ft.ElevatedButton("Orçamentos", on_click=show_orc, style=pill),
            ft.ElevatedButton("Cadastro de Clientes", on_click=show_cad, style=pill),
            ft.ElevatedButton("Gerar Contrato", on_click=show_contrato, style=pill),
            ft.ElevatedButton("Relatórios", on_click=show_relatorios, style=pill),
            ft.ElevatedButton("Usuários", on_click=show_usuarios, style=pill),
        ],
        spacing=12,
    )

    # ===================== Orçamentos =====================
    id_input = ft.TextField(label="ID Orçamento", width=260, disabled=True)

    def _copiar_id(e):
        if id_input.value:
            page.set_clipboard(id_input.value)
            page.snack_bar = ft.SnackBar(ft.Text("ID copiado para a área de transferência."), open=True)
            page.update()

    btn_copiar_id = ft.ElevatedButton("Copiar ID", on_click=_copiar_id, style=pill)

    tipo_servico_input = ft.Dropdown(
        label="Tipo de Serviço",
        options=[ft.dropdown.Option("Impressão"), ft.dropdown.Option("Digitalização")],
        width=220,
    )
    vendedor_input = ft.TextField(label="Vendedor", width=220)  # NOVO
    razao_input = ft.TextField(label="Razão Social/Nome", width=320)
    doc_tipo_orc = ft.Dropdown(
        label="Documento",
        options=[ft.dropdown.Option("CNPJ"), ft.dropdown.Option("CPF")],
        width=140,
        hint_text="Selecione o documento",
    )
    doc_input = ft.TextField(label="CNPJ/CPF", width=220, hint_text="Selecione o documento acima")
    # Botão Buscar ao lado do CNPJ/CPF
    orc_btn_buscar_CLIENTE = ft.ElevatedButton("Buscar CNPJ", style=pill)
    def _enable_orc_edit():
        _set_orc_editable(True)
        resultado_orc.value = "Edição liberada."
        page.update()
    btn_orc_editar = ft.ElevatedButton("Editar Orçamento", on_click=lambda e: require_admin(_enable_orc_edit), style=pill)

    email_input = ft.TextField(label="E-mail do CLIENTE", width=280, hint_text="CLIENTE@dominio.com.br")
    status_input = ft.Dropdown(
        label="Status do CLIENTE",
        options=[ft.dropdown.Option(x) for x in ["Sem desconto", "Com Desconto"]],
        width=160,
    )
    unidade_input = ft.Dropdown(
        label="Unidade de medida",
        options=[ft.dropdown.Option("Centímetros"), ft.dropdown.Option("Metro")],
        width=180,
        value="Centímetros",
    )
    quantidade_input = ft.TextField(label="Quantidade", width=140, hint_text="Ex: 1250")

    # Campos de desconto (Orçamento)
    desconto_qtd_input = ft.TextField(label="Período Desconto", width=180, hint_text="Ex.: 6")
    desconto_unid_input = ft.Dropdown(label="Unidade Desconto", width=180, options=[ft.dropdown.Option("Dias"), ft.dropdown.Option("Meses"), ft.dropdown.Option("Anos")])

    # Forma de Pagamento (Orçamento)
    forma_pgto_input = ft.TextField(label="Forma de Pagamento", width=260, hint_text="Ex.: PIX 30 dias")


    # Exibicao do periodo de desconto
    desconto_inicial_view = ft.TextField(label="Periodo Desc. Inicial", width=220, disabled=True)
    desconto_restante_view = ft.TextField(label="Periodo Desc. Faltante", width=230, disabled=True)



    resultado_orc = ft.Text("", size=13)

    # Busca/reImpressão
    orc_busca_id = ft.TextField(label="ID Orçamento", width=220)
    orc_busca_doc_tipo = ft.Dropdown(
        label="Documento",
        options=[ft.dropdown.Option("CNPJ"), ft.dropdown.Option("CPF")],
        width=140,
    )
    orc_busca_doc = ft.TextField(label="CNPJ/CPF", width=240)
    orc_tab_container = ft.Column(visible=False)

    # Estado de edição e helper para (des)habilitar campos do formulário de Orçamento
    orc_estado = {"editando": False, "salvo": True}

    def _set_orc_editable(enabled: bool):
        campos_text = [
            razao_input,
            vendedor_input,
            doc_input,
            email_input,
            quantidade_input,
            forma_pgto_input,
            desconto_qtd_input,
        ]
        campos_drop = [
            doc_tipo_orc,
            status_input,
            unidade_input,
            tipo_servico_input,
            desconto_unid_input,
        ]
        for c in campos_text + campos_drop:
            try:
                c.disabled = not enabled
            except Exception:
                pass
        try:
            orc_btn_buscar_CLIENTE.disabled = not enabled
        except Exception:
            pass
        id_input.disabled = True
        page.update()

    # Inicia bloqueado até clicar em "Novo Orcamento"
    _set_orc_editable(False)

    def _mask_orc_busca_doc(e):
        if orc_busca_doc_tipo.value == "CNPJ":
            orc_busca_doc.value = formatar_cnpj(orc_busca_doc.value or "")
        elif orc_busca_doc_tipo.value == "CPF":
            orc_busca_doc.value = formatar_cpf(orc_busca_doc.value or "")
        page.update()

    orc_busca_doc.on_change = _mask_orc_busca_doc
    def _dias_total_desc():
        try:
            n = int((desconto_qtd_input.value or "").strip())
        except Exception:
            return 0
        u = (desconto_unid_input.value or "").strip().lower()
        if n <= 0:
            return 0
        if u.startswith("dia"):
            return n
        if u.startswith("mes"):
            return n * 30
        if u.startswith("ano"):
            return n * 365
        return 0

    def _pluralize(n, base):
        if base == "Dia":
            return "Dia" if n == 1 else "Dias"
        if base == "Mes":
            return "Mes" if n == 1 else "Meses"
        if base == "Ano":
            return "Ano" if n == 1 else "Anos"
        return base

    def _atualizar_periodo_desc(e=None):
        qtxt = (desconto_qtd_input.value or "").strip()
        u = (desconto_unid_input.value or "").strip().lower()
        if qtxt and u:
            try:
                n = int(qtxt)
            except Exception:
                n = 0
            base = "Dia" if u.startswith("dia") else ("Mes" if u.startswith("mes") else ("Ano" if u.startswith("ano") else ""))
            desconto_inicial_view.value = (f"{qtxt} {_pluralize(n, base)}" if base else "")
        else:
            desconto_inicial_view.value = ""
        try:
            ref = None
            if doc_tipo_orc.value and (doc_input.value or "").strip():
                doc_fmt = formatar_doc(doc_tipo_orc.value, doc_input.value or "")
                ref = get_ultimo_pedido_data(doc_fmt)
            dias_tot = _dias_total_desc()
            if ref and dias_tot > 0:
                falt = max(0, (ref + timedelta(days=dias_tot) - datetime.now()).days)
                desconto_restante_view.value = f"{falt} Dia(s)"
            else:
                desconto_restante_view.value = ""
        except Exception:
            desconto_restante_view.value = ""
        page.update()

    def _fill_orc_form_from_dict(d):
        id_input.value = d.get("ID Orçamento") or ""
        tipo_servico_input.value = d.get("Tipo de Serviço") or None
        # Usa extração robusta do nome/Razão social do CLIENTE
        razao_input.value = extrair_nome_CLIENTE(d) or ""
        vendedor_input.value = d.get("Vendedor") or ""  # NOVO
        doc_tipo_orc.value = d.get("Documento") or None
        doc_input.value = d.get("CNPJ/CPF") or ""
        email_input.value = d.get("E-mail") or ""
        status_input.value = d.get("Status") or None
        unidade_input.value = d.get("Unidade") or None
        quantidade_input.value = d.get("Quantidade") or ""
        forma_pgto_input.value = d.get("Forma de Pagamento") or ""
        try:
            cad = None
            # 1) Primeiro tenta pegar do prprio Orçamento selecionado (variações toleradas)
            val_dur = str(_get_cad_field_val(d, "Desconto Duracao") or "")
            val_un  = str(_get_cad_field_val(d, "Desconto Unidade") or "")
            # 2) Se vazio, busca no cadastro do CLIENTE
            if not val_dur or not val_un:
                if doc_tipo_orc.value and (doc_input.value or "").strip():
                    cad = buscar_cadastro_por_documento(doc_tipo_orc.value, doc_input.value or "") or {}
                    if not val_dur:
                        val_dur = str(_get_cad_field_val(cad, "Desconto Duracao") or "")
                    if not val_un:
                        val_un = str(_get_cad_field_val(cad, "Desconto Unidade") or "")
            desconto_qtd_input.value = val_dur
            desconto_unid_input.value = (val_un or None)
        except Exception:
            pass
        _atualizar_periodo_desc()
        
        try:
            if doc_tipo_orc.value and (doc_input.value or "").strip():
                cad = buscar_cadastro_por_documento(doc_tipo_orc.value, doc_input.value or "") or {}
                desconto_qtd_input.value = str(_get_cad_field_val(cad, "Desconto Duracao") or "")
                desconto_unid_input.value = str(_get_cad_field_val(cad, "Desconto Unidade") or None)
        except Exception:
            pass
        _atualizar_periodo_desc()
        resultado_orc.value = "Orçamento carregado para reImpressão."
        page.update()

    def _render_tabela_generica(lista, on_select):
        orc_tab_container.controls.clear()
        if not lista:
            orc_tab_container.visible = False
            page.update()
            return
        cols = [
            ft.DataColumn(ft.Text("Selecionar")),
            ft.DataColumn(ft.Text("Vendedor")),
            ft.DataColumn(ft.Text("ID")),
            ft.DataColumn(ft.Text("CNPJ/CPF")),
            ft.DataColumn(ft.Text("Razão Social/Nome")),
            ft.DataColumn(ft.Text("Quantidade (m)")),
            ft.DataColumn(ft.Text("Valor Total")),
            ft.DataColumn(ft.Text("Data/Hora")),
        ]
        rows = []
        for d in lista:
            _id = str(d.get("ID Orçamento") or "")
            _doc = str(d.get("CNPJ/CPF") or "")
            # Usa extração robusta para exibir o nome do CLIENTE
            _nome = str(extrair_nome_CLIENTE(d) or "")
            _vendedor = str(d.get("Vendedor") or "")
            _metros = str(d.get("Metros") or "")
            _total = str(d.get("Valor Total") or "")
            _dh = str(d.get("Data/Hora") or "")
            btn_sel = ft.TextButton("Selecionar", on_click=lambda e, dd=d: on_select(dd))
            rows.append(
                ft.DataRow(
                    cells=[
                        ft.DataCell(btn_sel),
                        ft.DataCell(ft.Text(_vendedor)),
                        ft.DataCell(ft.Text(_id)),
                        ft.DataCell(ft.Text(_doc)),
                        ft.DataCell(ft.Text(_nome)),
                        ft.DataCell(ft.Text(_metros)),
                        ft.DataCell(ft.Text(_total)),
                        ft.DataCell(ft.Text(_dh)),
                    ]
                )
            )
        tabela = ft.DataTable(columns=cols, rows=rows, heading_row_height=32, data_row_min_height=32)
        orc_tab_container.controls.append(tabela)
        orc_tab_container.visible = True
        page.update()

    def buscar_orcamentos(e):
        idf = (orc_busca_id.value or "").strip()
        docf = None
        if orc_busca_doc_tipo.value and (orc_busca_doc.value or "").strip():
            if not validar_doc(orc_busca_doc_tipo.value, orc_busca_doc.value or ""):
                resultado_orc.value = f"{orc_busca_doc_tipo.value} inválido."
                orc_tab_container.visible = False
                orc_tab_container.controls.clear()
                page.update()
                return
            docf = formatar_doc(orc_busca_doc_tipo.value, orc_busca_doc.value or "")
        lista = get_orcamentos_list(doc_formatado=docf, id_orc=idf) if (idf or docf) else get_orcamentos_list()
        _render_tabela_generica(lista, on_select=_fill_orc_form_from_dict)

    def limpar_busca_orc(e):
        for c in [orc_busca_id, orc_busca_doc, resultado_orc]:
            c.value = ""
        orc_busca_doc_tipo.value = None
        orc_tab_container.controls.clear()
        orc_tab_container.visible = False
        page.update()

    def atualizar_hint_doc_orc(e=None):
        doc_input.hint_text = (
            "00.000.000/0000-00" if doc_tipo_orc.value == "CNPJ"
            else "000.000.000-00 (apenas RJ)" if doc_tipo_orc.value == "CPF"
            else "Selecione o documento acima"
        )
        page.update()

    doc_tipo_orc.on_change = lambda e: (setattr(doc_input, "value", ""), atualizar_hint_doc_orc(), _atualizar_periodo_desc())

    def aplicar_mascara_doc_orc(e):
        valor = doc_input.value or ""
        if doc_tipo_orc.value == "CNPJ":
            doc_input.value = formatar_cnpj(valor)
        elif doc_tipo_orc.value == "CPF":
            doc_input.value = formatar_cpf(valor)
        page.update()

    doc_input.on_change = lambda e: (aplicar_mascara_doc_orc(e), _atualizar_periodo_desc())
    def _get_cad_field_val(d: dict, target: str):
        try:
            def _norm(s: str) -> str:
                import re as _re
                return _re.sub(r'[^a-z]', '', (s or '').lower())
            t = _norm(target)
            for k in list(d.keys()):
                if _norm(str(k)) == t:
                    return d.get(k)
        except Exception:
            pass
        return ''

    def buscar_CLIENTE_orc(e):
        if not doc_tipo_orc.value or not (doc_input.value or "").strip():
            resultado_orc.value = "Selecione o tipo de documento e preencha o CNPJ/CPF."
            page.update()
            return
        if not validar_doc(doc_tipo_orc.value, doc_input.value or ""):
            resultado_orc.value = f"{doc_tipo_orc.value} inválido."
            page.update()
            return
        cad = buscar_cadastro_por_documento(doc_tipo_orc.value, doc_input.value or "")
        if not cad:
            resultado_orc.value = "CLIENTE sem cadastro."
            page.update()
            return
        # Preenche usando as possíveis colunas (compatével com planilhas antigas/novas)
        razao_input.value = extrair_nome_CLIENTE(cad)
        
        try:
            desconto_qtd_input.value = str(_get_cad_field_val(cad, 'Desconto Duracao') or '')
            desconto_unid_input.value = str(_get_cad_field_val(cad, 'Desconto Unidade') or None)
        except Exception:
            pass
        _atualizar_periodo_desc()
        resultado_orc.value = "Dados do CLIENTE carregados."
        page.update()

    orc_btn_buscar_CLIENTE.on_click = buscar_CLIENTE_orc

    quantidade_input.on_change = lambda e: (
        setattr(quantidade_input, "value", re.sub(r"[^0-9,\.]", "", (quantidade_input.value or "").strip())),
        page.update(),
    )
    tipo_servico_input.on_change = lambda e: (
        setattr(id_input, "value", gerar_id(tipo_servico_input.value) if tipo_servico_input.value else ""),
        page.update(),
    )
    unidade_input.on_change = lambda e: (
        setattr(quantidade_input, "hint_text", "Ex: 1250" if unidade_input.value == "Centímetros" else "Ex: 12,50"),
        page.update(),
    )
    atualizar_hint_doc_orc()
    desconto_qtd_input.on_change = _atualizar_periodo_desc
    desconto_unid_input.on_change = _atualizar_periodo_desc

    def novo_orcamento(e):
        orc_estado["editando"] = True
        orc_estado["salvo"] = False
        _set_orc_editable(True)
        for c in [razao_input, vendedor_input, doc_input, email_input, quantidade_input, forma_pgto_input]:
            c.value = ""
        for d in [doc_tipo_orc, status_input]:
            d.value = None
        try:
            unidade_input.value = unidade_input.value or "Centimetros"
        except Exception:
            pass
        tipo_servico_input.value = None
        id_input.disabled = True
        id_input.value = ""
        resultado_orc.value = "Novo orcamento iniciado."
        desconto_inicial_view.value = ""
        desconto_restante_view.value = ""
        page.update()

    def limpar_orcamento(e=None):
        # Limpa somente os campos do orcamento e destrava a tela
        for c in [razao_input, vendedor_input, doc_input, email_input, quantidade_input, forma_pgto_input]:
            c.value = ""
        for d in [doc_tipo_orc, status_input]:
            d.value = None
        desconto_qtd_input.value = ""
        desconto_unid_input.value = None
        desconto_inicial_view.value = ""
        desconto_restante_view.value = ""
        tipo_servico_input.value = None
        id_input.value = ""
        orc_estado["editando"] = False
        orc_estado["salvo"] = True
        _set_orc_editable(False)
        page.update()

    def salvar_orcamento(e):
        try:
            id_orc = id_input.value.strip()
            tipo_servico = tipo_servico_input.value
            CLIENTE_valor = razao_input.value.strip()
            vendedor = vendedor_input.value.strip()  # NOVO
            documento = doc_tipo_orc.value
            doc_val = doc_input.value.strip()
            email = email_input.value.strip()
            status = status_input.value
            unidade = unidade_input.value
            forma_pgto = (forma_pgto_input.value or "").strip()
            qtd = _parse_ptbr_float(quantidade_input.value)
            if not (CLIENTE_valor and documento and doc_val and email and status and id_orc and tipo_servico):
                resultado_orc.value = "Preencha todos os campos obrigatérios!"
                page.update()
                return
            if not validar_doc(documento, doc_val):
                resultado_orc.value = f"{documento} inválido!"
                page.update()
                return
            if not validar_email(email):
                resultado_orc.value = "E-mail inválido!"
                page.update()
                return
            CLIENTE_label = "Nome" if documento == "CPF" else "Razão Social"
            doc_fmt = formatar_doc(documento, doc_val)
            metros = qtd / 100.0 if unidade == "Centímetros" else qtd
            preco = 8.00 if status in ["Novo", "Ativo"] else 8.50
            valor_total = metros * preco
            dh = data_hora_tokens()
            salvar_excel_orcamento(
                {
                    "ID Orçamento": id_orc,
                    "Data/Hora": dh["combinado"],
                    "Tipo de Serviço": tipo_servico,
                    "CLIENTE (Etiqueta PDF)": CLIENTE_label,
                    "CLIENTE (Valor)": CLIENTE_valor,
                    "Documento": documento,
                    "CNPJ/CPF": doc_fmt,
                    "E-mail": email,
                    "Vendedor": vendedor,
                    "Status": status,
                    "Quantidade": format_num_ptbr(qtd),
                    "Unidade": unidade,
                    "Metros": format_num_ptbr(metros),
                    "Preço por metro": format_num_ptbr(preco),
                    "Forma de Pagamento": forma_pgto,
                    "Valor Total": format_num_ptbr(valor_total),
                }
            )
            resultado_orc.value = f"Orcamento salvo! Valor Total: R$ {format_num_ptbr(valor_total)}"
            page.update()
        except Exception as ex:
            resultado_orc.value = f"Erro: {ex}"
            page.update()

    def gerar_pdf_click(e):
        try:
            dialog = ft.FilePicker(on_result=lambda d: gerar_pdf_final(d))
            page.overlay.append(dialog)
            page.update()
            dialog.get_directory_path()
        except Exception as ex:
            resultado_orc.value = f"Erro: {ex}"
            page.update()

    def gerar_pdf_final(dlg_result):
        if not dlg_result.path:
            resultado_orc.value = "Nenhuma pasta selecionada."
            page.update()
            return
        try:
            id_orc = id_input.value.strip()
            tipo_servico = tipo_servico_input.value
            CLIENTE_valor = razao_input.value.strip()
            vendedor = vendedor_input.value.strip()
            documento = doc_tipo_orc.value
            doc_val = doc_input.value.strip()
            email = email_input.value.strip()
            status = status_input.value
            unidade = unidade_input.value
            forma_pgto = (forma_pgto_input.value or "").strip()
            qtd = _parse_ptbr_float(quantidade_input.value)
            if not validar_doc(documento, doc_val):
                resultado_orc.value = f"{documento} invalido!"
                page.update()
                return
            if not validar_email(email):
                resultado_orc.value = "E-mail invalido!"
                page.update()
                return
            CLIENTE_label = "Nome" if documento == "CPF" else "Razão Social"
            doc_fmt = formatar_doc(documento, doc_val)
            metros = qtd / 100.0 if unidade == "Centímetros" else qtd
            preco = 8.00 if status in ["Novo", "Ativo"] else 8.50
            valor_total = metros * preco
            dh = data_hora_tokens()
            caminho = gerar_pdf_orcamento(
                [
                    id_orc, dh["combinado"], tipo_servico, CLIENTE_label, CLIENTE_valor,
                    documento, doc_fmt, email, vendedor, status,
                    format_num_ptbr(qtd), unidade, format_num_ptbr(metros),
                    format_num_ptbr(preco),
                    forma_pgto,  # no PDF
                    format_num_ptbr(valor_total),
                ],
                dlg_result.path,
            )
            resultado_orc.value = f"PDF salvo em: {caminho}"
            page.update()
        except Exception as ex:
            resultado_orc.value = f"Erro: {ex}"
            page.update()

    campos_orc_row = ft.Row(
        [
            tipo_servico_input,
            vendedor_input,
            razao_input,
            doc_tipo_orc,
            doc_input,
            email_input,
            status_input,
            desconto_qtd_input,
            desconto_unid_input,
            unidade_input,
            quantidade_input,
            forma_pgto_input,
                    desconto_inicial_view,
            desconto_restante_view,
        ],
        wrap=True,
        spacing=10,
        run_spacing=8,
    )

    cont_orcamento = ft.Column(
        controls=[
            ft.Text("Fashion Tech - Audaces RJ e ES - Orçamento de Impressão de Riscos", size=18, weight="bold"),
            ft.Row([id_input], spacing=10, wrap=True),
            campos_orc_row,
            ft.Row(
                [
                    ft.ElevatedButton("Novo Orcamento", on_click=novo_orcamento, style=pill),
                    ft.ElevatedButton("Salvar Orcamento", on_click=salvar_orcamento, style=pill),
                    btn_orc_editar,
                    ft.ElevatedButton("Limpar Orcamento", on_click=limpar_orcamento, style=pill),
                    ft.ElevatedButton("Gerar PDF", on_click=gerar_pdf_click, style=pill),
                    btn_copiar_id,
                    orc_btn_buscar_CLIENTE,
                ],
                spacing=10,
                wrap=True,
            ),
            ft.Divider(),
            ft.Text("Buscar/Reimprimir Orçamento", weight="bold"),
            ft.Row(
                [
                    orc_busca_id,
                    orc_busca_doc_tipo,
                    orc_busca_doc,
                    ft.ElevatedButton("Buscar Orçamentos", on_click=buscar_orcamentos, style=pill),
                    ft.ElevatedButton("Limpar Pesquisa", on_click=limpar_busca_orc, style=pill),
                ],
                wrap=True,
                spacing=10,
            ),
            orc_tab_container,
            resultado_orc,
        ],
        visible=True,
    )

    # ===================== CADASTRO DE CLIENTES =====================
    # A) Documento (ok)
    cad_doc_tipo = ft.Dropdown(
        label="Documento",
        options=[ft.dropdown.Option("CNPJ"), ft.dropdown.Option("CPF")],
        width=220,
        hint_text="Selecione o documento",
    )
    cad_doc = ft.TextField(label="CNPJ/CPF", width=270, hint_text="Selecione o documento acima")
    btn_buscar_doc = ft.ElevatedButton("Buscar CNPJ/CPF", style=pill)
    def _enable_cad_edit():
        for w in [cad_doc_tipo, cad_doc, cad_razao, cad_nome_fantasia, cad_ie, cad_ie_status, cad_im, cad_im_status,
                  cad_situacao_cnpj, cad_contato, cad_tel1, cad_tel2, cad_email_cnpj, cad_email_manual, cad_cep,
                  cad_end, cad_num, cad_comp, cad_bairro, cad_municipio, cad_uf, cad_cep_entrega, cad_end_entrega,
                  cad_num_entrega, cad_comp_entrega, cad_bairro_entrega, cad_municipio_entrega, cad_uf_entrega]:
            try:
                w.disabled = False
            except Exception:
                pass
        cad_resultado.value = "Edição liberada."; page.update()
    btn_cad_editar = ft.ElevatedButton("Editar Cadastro", on_click=lambda e: require_admin(_enable_cad_edit), style=pill)

    # B) Informações
    cad_razao = ft.TextField(label="Razão Social/Nome", width=420)
    cad_nome_fantasia = ft.TextField(label="Nome Fantasia", width=420)
    cad_ie = ft.TextField(label="Inscrição Estadual", width=210)
    cad_ie_status = ft.TextField(label="Situação IE", width=210)
    cad_im = ft.TextField(label="Inscrição Municipal", width=210)
    cad_im_status = ft.TextField(label="Situação IM", width=210)
    cad_situacao_cnpj = ft.TextField(label="Situação Cadastral", width=220)

    # C) Contatos
    cad_contato = ft.TextField(label="Contato", width=280)
    cad_tel1 = ft.TextField(label="Telefone 1", width=180)
    cad_tel2 = ft.TextField(label="Telefone 2", width=180)
    cad_email_cnpj = ft.TextField(label="E-mail (CNPJ)", width=300, disabled=True)
    cad_email_manual = ft.TextField(label="E-mail (Manual)", width=300)

    # D) Endereço CNPJ/CPF (ok)
    cad_cep = ft.TextField(label="CEP", width=160, hint_text="00000-000")
    cad_end = ft.TextField(label="Endereço", width=420)
    cad_num = ft.TextField(label="Número", width=120)
    cad_comp = ft.TextField(label="Complemento", width=200)
    cad_bairro = ft.TextField(label="Bairro", width=220)
    cad_municipio = ft.TextField(label="Município", width=260)
    cad_uf = ft.TextField(label="UF", width=80, max_length=2)

    # E) Endereço de Entrega (ok)
    cad_cep_entrega = ft.TextField(label="CEP", width=160, hint_text="00000-000")
    cad_end_entrega = ft.TextField(label="Endereço", width=420)
    cad_num_entrega = ft.TextField(label="Número", width=120)
    cad_comp_entrega = ft.TextField(label="Complemento", width=200)
    cad_bairro_entrega = ft.TextField(label="Bairro", width=220)
    cad_municipio_entrega = ft.TextField(label="Município", width=260)
    cad_uf_entrega = ft.TextField(label="UF", width=80, max_length=2)

    # Campos automáticos bloqueados por padrão
    campos_auto = [
        cad_razao, cad_nome_fantasia, cad_end, cad_bairro, cad_municipio, cad_uf,
        cad_cep, cad_email_cnpj, cad_num, cad_comp, cad_ie, cad_ie_status,
        cad_im, cad_im_status, cad_situacao_cnpj,
    ]
    for c in campos_auto:
        c.disabled = True

    cad_resultado = ft.Text("", size=13)

    def _update_hint_cad_doc(e=None):
        cad_doc.hint_text = (
            "00.000.000/0000-00" if cad_doc_tipo.value == "CNPJ"
            else "000.000.000-00 (RJ)" if cad_doc_tipo.value == "CPF"
            else "Selecione o documento"
        )
        page.update()

    cad_doc_tipo.on_change = lambda e: (setattr(cad_doc, "value", ""), _update_hint_cad_doc())
    cad_doc.on_change = lambda e: (
        setattr(
            cad_doc,
            "value",
            formatar_cnpj(cad_doc.value) if cad_doc_tipo.value == "CNPJ"
            else formatar_cpf(cad_doc.value) if cad_doc_tipo.value == "CPF"
            else cad_doc.value,
        ),
        page.update(),
    )

    def _endereco_ok():
        return bool((cad_cep.value or "").strip()) and bool((cad_end.value or "").strip())

    def _atualizar_estado_copiar():
        btn_copiar_endereco.disabled = not _endereco_ok()

    cad_cep.on_change = lambda e: (
        setattr(cad_cep, "value", formatar_cep(cad_cep.value)),
        _atualizar_estado_copiar(),
        page.update(),
    )
    cad_cep_entrega.on_change = lambda e: (
        setattr(cad_cep_entrega, "value", formatar_cep(cad_cep_entrega.value)),
        page.update(),
    )

    for campo in [cad_end, cad_num, cad_comp, cad_bairro, cad_municipio, cad_uf]:
        campo.on_change = lambda e: (_atualizar_estado_copiar(), page.update())

    def copiar_endereco_cnpj(e):
        if not _endereco_ok():
            cad_resultado.value = "Preencha o Endereço principal antes de copiar."
            page.update()
            return
        cad_cep_entrega.value = cad_cep.value
        cad_end_entrega.value = cad_end.value
        cad_num_entrega.value = cad_num.value
        cad_comp_entrega.value = cad_comp.value
        cad_bairro_entrega.value = cad_bairro.value
        cad_municipio_entrega.value = cad_municipio.value
        cad_uf_entrega.value = cad_uf.value
        cad_resultado.value = "Endereço de Entrega copiado."
        page.update()

    def tentar_busca_cep(e):
        cep = re.sub(r"\D", "", cad_cep.value or "")
        if len(cep) != 8:
            cad_resultado.value = "CEP inválido."
            page.update()
            return
        try:
            data = http_get_json(f"https://viacep.com.br/ws/{cep}/json/", timeout=8)
            if data.get("erro"):
                cad_resultado.value = "CEP não encontrado."
            else:
                cad_end.value = data.get("LOGRADOURO", "") or cad_end.value
                cad_bairro.value = data.get("bairro", "") or cad_bairro.value
                cad_municipio.value = data.get("localidade", "") or cad_municipio.value
                cad_uf.value = data.get("uf", "") or cad_uf.value
                cad_resultado.value = "Endereço preenchido pelo CEP (confira os campos)."
        except Exception as ex:
            cad_resultado.value = f"Consulta de CEP indisponível ({ex}). Preencha manualmente."
        _atualizar_estado_copiar()
        page.update()

    def tentar_busca_cep_entrega(e):
        cep = re.sub(r"\D", "", cad_cep_entrega.value or "")
        if len(cep) != 8:
            cad_resultado.value = "CEP inválido (Entrega)."
            page.update()
            return
        try:
            data = http_get_json(f"https://viacep.com.br/ws/{cep}/json/", timeout=8)
            if data.get("erro"):
                cad_resultado.value = "CEP de entrega não encontrado."
            else:
                cad_end_entrega.value = data.get("LOGRADOURO", "") or cad_end_entrega.value
                cad_bairro_entrega.value = data.get("bairro", "") or cad_bairro_entrega.value
                cad_municipio_entrega.value = data.get("localidade", "") or cad_municipio_entrega.value
                cad_uf_entrega.value = data.get("uf", "") or cad_uf_entrega.value
                cad_resultado.value = "Endereço de entrega preenchido pelo CEP."
        except Exception as ex:
            cad_resultado.value = f"Consulta de CEP (entrega) indisponível ({ex}). Preencha manualmente."
        page.update()

    def _extrair_do_brasilapi(data: dict) -> dict:
        out = {}
        out["razao"] = (data.get("razao_social") or "").strip()
        out["fantasia"] = (data.get("nome_fantasia") or "").strip()

        est = data.get("estabelecimento") or {}
        out["LOGRADOURO"] = (est.get("LOGRADOURO") or data.get("LOGRADOURO") or "").strip()
        out["numero"] = str(est.get("numero") or data.get("numero") or "")
        out["complemento"] = str(est.get("complemento") or data.get("complemento") or "")
        out["bairro"] = (est.get("bairro") or data.get("bairro") or "").strip()
        out["municipio"] = (est.get("cidade") or est.get("municipio") or data.get("municipio") or "").strip()
        out["uf"] = (est.get("estado") or est.get("uf") or data.get("uf") or "").strip()
        out["cep"] = (est.get("cep") or data.get("cep") or "")
        out["email"] = (est.get("email") or data.get("email") or "")

        out["situacao_cadastral"] = (
            data.get("descricao_situacao_cadastral") or data.get("situacao_cadastral") or est.get("situacao_cadastral") or ""
        ).strip()

        ie, ie_status, im, im_status = "", "", "", ""
        ies = est.get("inscricoes_estaduais") or data.get("inscricoes_estaduais")
        if isinstance(ies, list) and ies:
            item = ies[0]
            ie = str(item.get("inscricao_estadual") or item.get("ie") or "").strip()
            if "situacao" in item and item["situacao"]:
                ie_status = str(item.get("situacao"))
            elif "ativo" in item:
                ie_status = "Habilitado" if item.get("ativo") else "Inativo"
        else:
            ie = str(est.get("inscricao_estadual") or data.get("inscricao_estadual") or "").strip()
            ie_status = str(est.get("situacao_ie") or data.get("situacao_ie") or "").strip()

        im = str(est.get("inscricao_municipal") or data.get("inscricao_municipal") or "").strip()
        im_status = str(est.get("situacao_im") or data.get("situucao_im") or "").strip()

        out["ie"], out["ie_status"], out["im"], out["im_status"] = ie, ie_status, im, im_status
        return out

    def _preencher_campos_from(d):
        cad_razao.value = d.get("razao", "") or cad_razao.value
        cad_nome_fantasia.value = d.get("fantasia", "") or cad_nome_fantasia.value
        cad_end.value = d.get("LOGRADOURO", "") or cad_end.value
        cad_num.value = d.get("numero", "") or cad_num.value
        cad_comp.value = d.get("complemento", "") or cad_comp.value
        cad_bairro.value = d.get("bairro", "") or cad_bairro.value
        cad_municipio.value = d.get("municipio", "") or cad_municipio.value
        uf = d.get("uf", "")
        cad_uf.value = uf or cad_uf.value
        cep = d.get("cep", "")
        cad_cep.value = formatar_cep(cep) if cep else cad_cep.value
        em = d.get("email", "")
        cad_email_cnpj.value = em or cad_email_cnpj.value
        if d.get("ie"):
            cad_ie.value = d.get("ie")
        if d.get("ie_status"):
            cad_ie_status.value = d.get("ie_status")
        if d.get("im"):
            cad_im.value = d.get("im")
        if d.get("im_status"):
            cad_im_status.value = d.get("im_status")
        sc = d.get("situacao_cadastral", "")
        cad_situacao_cnpj.value = sc or cad_situacao_cnpj.value

    def liberar_campos_auto():
        for c in campos_auto:
            c.disabled = False

    def buscar_geral(e):
        ja = buscar_cadastro_por_documento(cad_doc_tipo.value or "", cad_doc.value or "") if cad_doc_tipo.value else None
        if ja:
            for k, v in {
                cad_razao: "Razão Social/Nome",
                cad_nome_fantasia: "Nome Fantasia",
                cad_contato: "Contato",
                cad_ie: "Inscrição Estadual",
                cad_ie_status: "Situação IE",
                cad_im: "Inscrição Municipal",
                cad_im_status: "Situação IM",
                cad_situacao_cnpj: "Situação Cadastral",
                cad_cep: "CEP",
                cad_end: "Endereço",
                cad_num: "Número",
                cad_comp: "Complemento",
                cad_bairro: "Bairro",
                cad_municipio: "Município",
                cad_uf: "UF",
                cad_tel1: "Telefone 1",
                cad_tel2: "Telefone 2",
                cad_email_cnpj: "E-mail (CNPJ)",
                cad_email_manual: "E-mail (Manual)",
                cad_cep_entrega: "Entrega CEP",
                cad_end_entrega: "Entrega Endereço",
                cad_num_entrega: "Entrega Número",
                cad_comp_entrega: "Entrega Complemento",
                cad_bairro_entrega: "Entrega Bairro",
                cad_municipio_entrega: "Entrega Município",
                cad_uf_entrega: "Entrega UF",
            }.items():
                k.value = str(ja.get(v) or "")
            cad_resultado.value = "Documento já cadastrado. Dados carregados."
            _atualizar_estado_copiar()
            page.update()
            return

        if cad_doc_tipo.value == "CNPJ":
            cnpj = re.sub(r"\D", "", cad_doc.value or "")
            if len(cnpj) != 14 or not validar_cnpj(cad_doc.value or ""):
                cad_resultado.value = "CNPJ inválido."
                liberar_campos_auto()
                page.update()
                return
            ok = False
            try:
                d1 = http_get_json(f"https://brasilapi.com.br/api/cnpj/v1/{cnpj}", timeout=10)
                _preencher_campos_from(_extrair_do_brasilapi(d1))
                ok = True
            except Exception:
                pass
            if ok:
                cad_resultado.value = "Dados preenchidos (confira e complete)."
            else:
                cad_resultado.value = "não foi possível consultar o CNPJ agora. Preencha manualmente."
                liberar_campos_auto()
            _atualizar_estado_copiar()
            page.update()
        else:
            cad_resultado.value = "Consulta externa de CPF indisponível. Preencha manualmente."
            liberar_campos_auto()
            _atualizar_estado_copiar()
            page.update()

    def salvar_cadastro_click(e):
        try:
            if not cad_doc_tipo.value or not validar_doc(cad_doc_tipo.value, cad_doc.value or ""):
                cad_resultado.value = f"{cad_doc_tipo.value or 'Documento'} inválido."
                page.update()
                return
            email_cnpj_val = (cad_email_cnpj.value or "").strip()
            email_manual_val = (cad_email_manual.value or "").strip()
            if email_cnpj_val and not validar_email(email_cnpj_val):
                cad_resultado.value = "E-mail (CNPJ) inválido."
                page.update()
                return
            if email_manual_val and not validar_email(email_manual_val):
                cad_resultado.value = "E-mail (Manual) inválido."
                page.update()
                return

            dados = {
                "Documento": cad_doc_tipo.value,
                "CNPJ/CPF": formatar_doc(cad_doc_tipo.value, cad_doc.value or ""),
                "Razão Social/Nome": (cad_razao.value or "").strip(),
                "Nome Fantasia": (cad_nome_fantasia.value or "").strip(),
                "Contato": (cad_contato.value or "").strip(),
                "Inscrição Estadual": (cad_ie.value or "").strip(),
                "Situação IE": (cad_ie_status.value or "").strip(),
                "Inscrição Municipal": (cad_im.value or "").strip(),
                "Situação IM": (cad_im_status.value or "").strip(),
                "Situação Cadastral": (cad_situacao_cnpj.value or "").strip(),
                "CEP": formatar_cep(cad_cep.value or ""),
                "Endereço": (cad_end.value or "").strip(),
                "Número": (cad_num.value or "").strip(),
                "Complemento": (cad_comp.value or "").strip(),
                "Bairro": (cad_bairro.value or "").strip(),
                "Município": (cad_municipio.value or "").strip(),
                "UF": (cad_uf.value or "").strip().upper(),
                "Telefone 1": (cad_tel1.value or "").strip(),
                "Telefone 2": (cad_tel2.value or "").strip(),
                "E-mail (CNPJ)": email_cnpj_val,
                "E-mail (Manual)": email_manual_val,
                "Entrega CEP": formatar_cep(cad_cep_entrega.value or ""),
                "Entrega Endereço": (cad_end_entrega.value or "").strip(),
                "Entrega Número": (cad_num_entrega.value or "").strip(),
                "Entrega Complemento": (cad_comp_entrega.value or "").strip(),
                "Entrega Bairro": (cad_bairro_entrega.value or "").strip(),
                "Entrega Município": (cad_municipio_entrega.value or "").strip(),
                "Entrega UF": (cad_uf_entrega.value or "").strip().upper(),
            }

            ja = buscar_cadastro_por_documento(cad_doc_tipo.value, cad_doc.value or "")
            if ja:
                ok = atualizar_excel_cadastro(dados["CNPJ/CPF"], dados)
                cad_resultado.value = "Cadastro atualizado." if ok else "não foi possível atualizar."
            else:
                salvar_excel_cadastro(dados)
                cad_resultado.value = "Cadastro salvo."
            page.update()
        except Exception as ex:
            cad_resultado.value = f"Erro: {ex}"
            page.update()

    def limpar_cadastro(e):
        for c in [
            cad_doc_tipo, cad_doc, cad_razao, cad_nome_fantasia, cad_contato, cad_ie, cad_ie_status,
            cad_im, cad_im_status, cad_situacao_cnpj, cad_end, cad_num, cad_comp, cad_cep,
            cad_bairro, cad_municipio, cad_uf, cad_tel1, cad_tel2, cad_email_cnpj, cad_email_manual,
            cad_cep_entrega, cad_end_entrega, cad_num_entrega, cad_comp_entrega, cad_bairro_entrega,
            cad_municipio_entrega, cad_uf_entrega,
        ]:
            if hasattr(c, "value"):
                c.value = ""
        for c in campos_auto:
            c.disabled = True
        cad_resultado.value = "formulário limpo."
        _atualizar_estado_copiar()
        _update_hint_cad_doc()
        page.update()

    senha_admin = "906835"
    pwd_ref = {"dlg": None, "senha": None, "erro": None}

    def _confirmar_senha(ev):
        if (pwd_ref["senha"].value or "").strip() == senha_admin:
            for c in campos_auto:
                c.disabled = False
            pwd_ref["dlg"].open = False
            cad_resultado.value = "Modo edição habilitado."
            page.update()
        else:
            pwd_ref["erro"].value = "Senha incorreta."
            page.update()

    def abrir_edicao(e):
        if not pwd_ref["dlg"]:
            pwd_ref["senha"] = ft.TextField(label="Senha", password=True, can_reveal_password=True, width=220)
            pwd_ref["erro"] = ft.Text("", color="red")
            dlg = ft.AlertDialog(
                modal=True,
                title=ft.Text("Editar Cadastro"),
                content=ft.Column([pwd_ref["senha"], pwd_ref["erro"]], tight=True, spacing=6),
                actions=[
                    ft.TextButton("Cancelar", on_click=lambda ev: (setattr(pwd_ref["dlg"], "open", False), page.update())),
                    ft.ElevatedButton("Confirmar", on_click=_confirmar_senha, style=pill),
                ],
            )
            pwd_ref["dlg"] = dlg
        else:
            pwd_ref["senha"].value = ""
            pwd_ref["erro"].value = ""
        page.dialog = pwd_ref["dlg"]
        pwd_ref["dlg"].open = True
        page.update()

    def bloquear_campos(e=None):
        for c in campos_auto:
            c.disabled = True
        cad_resultado.value = "Campos bloqueados novamente."
        page.update()

    btn_buscar_cep = ft.ElevatedButton("Buscar CEP", on_click=tentar_busca_cep, style=pill)
    btn_buscar_doc = ft.ElevatedButton("Buscar CNPJ", on_click=buscar_geral, style=pill)
    btn_buscar_cep_entrega = ft.ElevatedButton("Buscar CEP (Entrega)", on_click=tentar_busca_cep_entrega, style=pill)
    btn_copiar_endereco = ft.OutlinedButton("Copiar Endereço CNPJ", on_click=copiar_endereco_cnpj, style=pill)
    btn_copiar_endereco.disabled = True

    bloco_documento = ft.Column(
        [ft.Text("Documento", size=16, weight="bold"), ft.Row([cad_doc_tipo, cad_doc, btn_buscar_doc, btn_cad_editar], wrap=True, spacing=12)],
        spacing=8,
    )

    # B) Informações
    bloco_informacoes = ft.Column(
        [
            ft.Text("Informações", size=16, weight="bold"),
            ft.Row([cad_razao, cad_nome_fantasia], wrap=True, spacing=12),
            ft.Row([cad_ie, cad_ie_status, cad_im, cad_im_status, cad_situacao_cnpj], wrap=True, spacing=12),
        ],
        spacing=8,
    )

    # C) Contatos
    bloco_contatos = ft.Column(
        [
            ft.Text("Contatos", size=16, weight="bold"),
            ft.Row([cad_contato, cad_tel1, cad_tel2], wrap=True, spacing=12),
            ft.Row([cad_email_cnpj, cad_email_manual], wrap=True, spacing=12),
        ],
        spacing=8,
    )

    # D) Endereço CNPJ/CPF
    bloco_endereco = ft.Column(
        [
            ft.Text("Endereço CNPJ/CPF", weight="bold", size=16),
            ft.Row([cad_cep, btn_buscar_cep], wrap=True, spacing=12),
            ft.Row([cad_end, cad_num, cad_comp], wrap=True, spacing=12),
            ft.Row([cad_bairro, cad_municipio, cad_uf], wrap=True, spacing=12),
        ],
        spacing=8,
    )

    # E) Endereço de Entrega
    bloco_entrega = ft.Column(
        [
            ft.Row(
                [ft.Text("Endereço de Entrega", weight="bold", size=16), btn_copiar_endereco],
                spacing=12,
                alignment="start",
                vertical_alignment="center",
            ),
            ft.Row([cad_cep_entrega, btn_buscar_cep_entrega], wrap=True, spacing=12),
            ft.Row([cad_end_entrega, cad_num_entrega, cad_comp_entrega], wrap=True, spacing=12),
            ft.Row([cad_bairro_entrega, cad_municipio_entrega, cad_uf_entrega], wrap=True, spacing=12),
        ],
        spacing=8,
    )

    cont_cadastro = ft.Column(
        controls=[
            ft.Text("Cadastro de Clientes", size=20, weight="bold"),
            bloco_documento,
            ft.Divider(),
            bloco_informacoes,
            ft.Divider(),
            bloco_contatos,
            ft.Divider(),
            bloco_endereco,
            ft.Divider(),
            bloco_entrega,
            ft.Divider(),
            ft.Row(
                [
                    ft.ElevatedButton("Salvar Cadastro", on_click=salvar_cadastro_click, style=pill),
                    ft.ElevatedButton("Editar Cadastro", on_click=abrir_edicao, style=pill),
                    ft.ElevatedButton("Bloquear novamente", on_click=bloquear_campos, style=pill),
                    ft.ElevatedButton("Limpar", on_click=limpar_cadastro, style=pill),
                ],
                spacing=12,
                wrap=True,
            ),
            cad_resultado,
        ],
        visible=False,
    )

    # ===================== GERAR CONTRATO =====================
    contrato_id_orc = ft.TextField(label="ID Orçamento", width=220, hint_text="Ex.: OR-IM...")
    contrato_doc_tipo = ft.Dropdown(label="Documento", options=[ft.dropdown.Option("CNPJ"), ft.dropdown.Option("CPF")], width=180)
    contrato_doc = ft.TextField(label="CNPJ/CPF", width=200)
    contrato_forma_pg = ft.TextField(label="Forma de pagamento", width=240, hint_text="Ex.: PIX 30 dias")
    # Comissão do vendedor (padrão 5%)
    contrato_comissao_vendedor = ft.TextField(label="% Comissão Vendedor (padrão 5%)", width=300, hint_text="5%", value="5%")
    contrato_result = ft.Text("", size=13)

    tabela_container = ft.Column(visible=False)
    selecionado_id_ref = {"id": ""}

    def _mask_contrato_doc(e):
        if contrato_doc_tipo.value == "CNPJ":
            contrato_doc.value = formatar_cnpj(contrato_doc.value or "")
        elif contrato_doc_tipo.value == "CPF":
            contrato_doc.value = formatar_cpf(contrato_doc.value or "")
        page.update()

    contrato_doc.on_change = _mask_contrato_doc

    btn_buscar_contrato = ft.ElevatedButton("Buscar Orçamentos", style=pill)
    btn_limpar_contrato = ft.ElevatedButton("Limpar Pesquisa", style=pill)
    btn_pdf_contrato = ft.ElevatedButton("Gerar PDF", style=pill)

    def _selecionar_contrato_from_row(d_orc: dict):
        selecionado_id_ref["id"] = d_orc.get("ID Orçamento") or ""
        contrato_id_orc.value = selecionado_id_ref["id"]
        contrato_doc_tipo.value = d_orc.get("Documento") or None
        contrato_doc.value = d_orc.get("CNPJ/CPF") or ""
        contrato_id_orc.disabled = True
        contrato_doc_tipo.disabled = True
        contrato_doc.disabled = True
        btn_buscar_contrato.disabled = True
        contrato_result.value = f"ID selecionado: {selecionado_id_ref['id']}"
        page.update()

    def _render_tabela_contrato(lista: list[dict]):
        tabela_container.controls.clear()
        if not lista:
            tabela_container.visible = False
            page.update()
            return
        cols = [
            ft.DataColumn(ft.Text("Selecionar")),
            ft.DataColumn(ft.Text("Vendedor")),
            ft.DataColumn(ft.Text("ID")),
            ft.DataColumn(ft.Text("CNPJ/CPF")),
            ft.DataColumn(ft.Text("Razão Social/Nome")),
            ft.DataColumn(ft.Text("Quantidade (m)")),
            ft.DataColumn(ft.Text("Valor Total")),
            ft.DataColumn(ft.Text("Data/Hora")),
        ]
        rows = []
        def _calc_metros_total(dd: dict) -> tuple[str, str]:
            preco_num = _parse_ptbr_float(str(dd.get("Preço por metro") or "0"))
            # Metros preferencialmente do campo, senão calcula
            metros_txt = str(dd.get("Metros") or "")
            metros_num = _parse_ptbr_float(metros_txt)
            if not re.match(r"^\d{1,3}(?:\.\d{3})*,\d{1,3}$", metros_txt or ""):
                qtd_num = _parse_ptbr_float(str(dd.get("Quantidade") or "0"))
                uni = str(dd.get("Unidade") or "")
                if qtd_num > 0:
                    metros_num = (qtd_num / 100.0) if uni.lower().startswith("cent") else qtd_num
            metros_fmt = format_num_ptbr(metros_num) if metros_num > 0 else (metros_txt or "")

            # Valor total prioritariamente do campo, senão calcula
            vtotal_txt = str(dd.get("Valor Total") or "")
            if not re.match(r"^\d{1,3}(?:\.\d{3})*,\d{2}$", vtotal_txt or ""):
                if metros_num <= 0:
                    metros_num = _parse_ptbr_float(str(dd.get("Metros") or "0"))
                if metros_num > 0 and preco_num > 0:
                    vtotal_txt = format_num_ptbr(metros_num * preco_num)
            return metros_fmt, vtotal_txt
        for d in lista:
            btn_sel = ft.TextButton("Selecionar", on_click=lambda e, dd=d: _selecionar_contrato_from_row(dd))
            metros_fmt, vtotal_fmt = _calc_metros_total(d)
            rows.append(
                ft.DataRow(
                    cells=[
                        ft.DataCell(btn_sel),
                        ft.DataCell(ft.Text(str(d.get("Vendedor") or ""))),
                        ft.DataCell(ft.Text(str(d.get("ID Orçamento") or ""))),
                        ft.DataCell(ft.Text(str(d.get("CNPJ/CPF") or ""))),
                        ft.DataCell(ft.Text(extrair_nome_CLIENTE(d))),
                        ft.DataCell(ft.Text(metros_fmt)),
                        ft.DataCell(ft.Text(vtotal_fmt)),
                        ft.DataCell(ft.Text(str(d.get("Data/Hora") or ""))),
                    ]
                )
            )
        tabela = ft.DataTable(columns=cols, rows=rows, heading_row_height=32, data_row_min_height=32)
        tabela_container.controls.append(tabela)
        tabela_container.visible = True
        page.update()

    def buscar_por_campos(e):
        selecionado_id_ref["id"] = ""
        contrato_result.value = ""
        idf = (contrato_id_orc.value or "").strip()
        docf = None
        if contrato_doc_tipo.value and (contrato_doc.value or "").strip():
            if not validar_doc(contrato_doc_tipo.value, contrato_doc.value or ""):
                contrato_result.value = f"{contrato_doc_tipo.value} inválido."
                tabela_container.visible = False
                tabela_container.controls.clear()
                page.update()
                return
            docf = formatar_doc(contrato_doc_tipo.value, contrato_doc.value or "")
        lista = get_orcamentos_list() if not (idf or docf) else get_orcamentos_list(doc_formatado=docf, id_orc=idf)
        contrato_result.value = f"{len(lista)} Orçamento(s) encontrados."
        _render_tabela_contrato(lista)

    def limpar_pesquisa(e):
        selecionado_id_ref["id"] = ""
        contrato_id_orc.value = ""
        contrato_doc_tipo.value = None
        contrato_doc.value = ""
        contrato_id_orc.disabled = False
        contrato_doc_tipo.disabled = False
        contrato_doc.disabled = False
        btn_buscar_contrato.disabled = False
        tabela_container.controls.clear()
        tabela_container.visible = False
        contrato_result.value = "Pesquisa limpa."
        page.update()

    btn_buscar_contrato.on_click = buscar_por_campos
    btn_limpar_contrato.on_click = limpar_pesquisa

    def _montar_contexto_contrato(d_orc):
        documento = str(d_orc.get("Documento") or "")
        doc_valor = str(d_orc.get("CNPJ/CPF") or "")
        # Usa extração robusta do nome/Razão social do CLIENTE
        CLIENTE_nome = extrair_nome_CLIENTE(d_orc)
        email = str(d_orc.get("E-mail") or "")

        cad = buscar_cadastro_por_documento(documento, doc_valor) or {}
        end_entrega_fmt = montar_endereco_entrega_formatado(cad)
        telefone = str(cad.get("Telefone 1") or cad.get("Telefone 2") or "").strip()

        tipo_servico = str(d_orc.get("Tipo de Serviço") or "")
        metros = str(d_orc.get("Metros") or "")
        valor_unit = str(d_orc.get("Preço por metro") or "")
        valor_total = str(d_orc.get("Valor Total") or "")

        empresa_endereco_concat = montar_endereco_entrega_formatado(
            {
                "Endereço": cad.get("Endereço"),
                "Número": cad.get("Número"),
                "Complemento": cad.get("Complemento"),
                "Bairro": cad.get("Bairro"),
                "Município": cad.get("Município"),
                "UF": cad.get("UF"),
                "CEP": cad.get("CEP"),
            }
        )

        return {
            "id_orc": d_orc["ID Orçamento"],
            "CLIENTE": CLIENTE_nome,
            "doc_valor": doc_valor,
            "email": email,
            "telefone": telefone,
            "end_entrega": end_entrega_fmt,
            "empresa_razao": CLIENTE_nome,
            "empresa_cnpj": doc_valor,
            "empresa_endereco_concat": empresa_endereco_concat,
            "forma_pgto": contrato_forma_pg.value or "",
            "tipo_servico": tipo_servico,
            "total_metros": metros,
            "valor_unit": valor_unit,
            "valor_total": valor_total,
        }

    def _salvar_contrato_core(d_orc, pasta, to_pdf: bool):
        ctx = _montar_contexto_contrato(d_orc)
        caminho_docx, err = gerar_contrato_docx(ctx, pasta)
        if err:
            contrato_result.value = err
            page.update()
            return
        metros = d_orc.get("Metros") or "0,00"
        pedido_num = get_proximo_pedido_numero()
        dh = data_hora_tokens()["combinado"]
        # Comissão: vendedor (default 5%), ADM (fixo 1%)
        def _parse_pct(txt: str | None) -> float:
            t = (txt or "").strip().replace("%", "").replace(",", ".")
            try:
                return float(t)
            except Exception:
                return 5.0
        pct_vendedor = _parse_pct(contrato_comissao_vendedor.value) if (contrato_comissao_vendedor.value or "").strip() else 5.0
        pct_adm = 1.0
        valor_total_float = _parse_ptbr_float(d_orc.get("Valor Total") or "0")
        valor_comissao_vendedor = valor_total_float * (pct_vendedor / 100.0)
        valor_comissao_adm = valor_total_float * (pct_adm / 100.0)
        vendedor_nome = d_orc.get("Vendedor") or ""
        salvar_excel_pedido(
            {
                "ID": f"CT-{pedido_num}",
                "Pedido": pedido_num,
                "Tipo de Serviço": d_orc.get("Tipo de Serviço") or "",
                "Status do CLIENTE": d_orc.get("Status") or "",
                "Quantidade (m)": metros,
                "Valor Unitário": d_orc.get("Preço por metro") or "",
                "Valor Total": d_orc.get("Valor Total") or "",
                "Data/Hora da criação do pedido": dh,
                "ID Orçamento": d_orc.get("ID Orçamento") or "",
                "Documento": d_orc.get("Documento") or "",
                "CNPJ/CPF": d_orc.get("CNPJ/CPF") or "",
                "CLIENTE": extrair_nome_CLIENTE(d_orc) or "",
                "Vendedor": vendedor_nome,
                "Forma de Pagamento Orçamento": d_orc.get("Forma de Pagamento") or "",     # NOVO
                "Forma de Pagamento Contrato": contrato_forma_pg.value or "",             # NOVO
                "% Comissão Vendedor": f"{str(pct_vendedor).replace('.',',')}%",
                "Valor Comissão Vendedor": f"R$ {format_num_ptbr(valor_comissao_vendedor)}",
                "% Comissão ADM": "1%",
                "Valor Comissão ADM": f"R$ {format_num_ptbr(valor_comissao_adm)}",
            }
        )
        if to_pdf:
            caminho_pdf, errp = converter_contrato_para_pdf(caminho_docx, pasta)
            contrato_result.value = errp or f"Contrato salvo (DOCX/PDF): {caminho_pdf}"
        else:
            contrato_result.value = f"Contrato salvo (DOCX): {caminho_docx}"
        page.update()

    def gerar_contrato_pdf_click(e):
        try:
            alvo_id = selecionado_id_ref["id"] or (contrato_id_orc.value or "").strip()
            if not alvo_id:
                contrato_result.value = "Selecione uma linha na tabela ou informe um ID Orçamento."
                page.update()
                return
            d = get_orcamento_by_id(alvo_id)
            if not d:
                contrato_result.value = "Orçamento não encontrado."
                page.update()
                return

            def on_dir(res):
                try:
                    if not res.path:
                        contrato_result.value = "Operação cancelada."
                        page.update()
                        return
                    _salvar_contrato_core(d, res.path, True)
                except Exception as ex:
                    contrato_result.value = f"Erro ao salvar contrato: {ex}"
                    page.update()

            dlg = ft.FilePicker(on_result=on_dir)
            page.overlay.append(dlg)
            page.update()
            dlg.get_directory_path()
        except Exception as ex:
            contrato_result.value = f"Erro: {ex}"
            page.update()

    btn_pdf_contrato.on_click = gerar_contrato_pdf_click

    linha_campos_contrato = ft.Row(
        [contrato_id_orc, contrato_doc_tipo, contrato_doc, contrato_forma_pg, contrato_comissao_vendedor],
        wrap=True,
        spacing=10,
        run_spacing=8,
    )

    cont_contrato = ft.Column(
        controls=[
            ft.Text("Gerar Contrato de Impressão", size=18, weight="bold"),
            linha_campos_contrato,
            ft.Row([
                btn_buscar_contrato,
                btn_limpar_contrato,
                btn_pdf_contrato,
                ft.ElevatedButton(
                    "Editar Contrato",
                    on_click=lambda e: require_admin(
                        lambda: (
                            setattr(contrato_id_orc, 'disabled', False),
                            setattr(contrato_doc_tipo, 'disabled', False),
                            setattr(contrato_doc, 'disabled', False),
                            setattr(btn_buscar_contrato, 'disabled', False),
                            page.update()
                        )
                    ),
                    style=pill,
                ),
            ], spacing=10, wrap=True),
            tabela_container,
            contrato_result,
        ],
        visible=False,
    )

    # ===================== Relatórios =====================
    def _fetch_all(endpoint: str) -> list[dict]:
        try:
            r = api_get(endpoint)
            if isinstance(r, dict) and "rows" in r:
                return r.get("rows") or []
            if isinstance(r, list):
                return r
        except Exception:
            pass
        return []

    def _exportar_para_excel(pasta: str):
        try:
            from openpyxl import Workbook
        except Exception as ex:
            resultado_global.value = f"openpyxl não disponível: {ex}"
            page.update(); return
        orcs = _fetch_all('/api/orcamentos')
        cads = _fetch_all('/api/cadastros')
        peds = _fetch_all('/api/pedidos')
        wb = Workbook(); ws = wb.active; ws.title = 'Orcamentos'
        if orcs:
            cols = list(orcs[0].keys()); ws.append(cols)
            for r in orcs: ws.append([r.get(k, '') for k in cols])
        ws2 = wb.create_sheet('Cadastros')
        if cads:
            cols = list(cads[0].keys()); ws2.append(cols)
            for r in cads: ws2.append([r.get(k, '') for k in cols])
        ws3 = wb.create_sheet('Pedidos')
        if peds:
            cols = list(peds[0].keys()); ws3.append(cols)
            for r in peds: ws3.append([r.get(k, '') for k in cols])
        nome = os.path.join(pasta, f"Relatorios_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
        wb.save(nome)
        resultado_global.value = f"Relatório gerado: {nome}"
        page.update()

    _dlg_rel = ft.FilePicker(on_result=lambda d: (_exportar_para_excel(d.path) if d.path else None))
    page.overlay.append(_dlg_rel)

    rel_tipo = ft.Dropdown(label='Tipo', options=[ft.dropdown.Option('Tudo'), ft.dropdown.Option('Orçamentos'), ft.dropdown.Option('Cadastros'), ft.dropdown.Option('Pedidos')], value='Tudo', width=200)

    def _exportar_por_tipo(pasta: str, tipo: str):
        try:
            from openpyxl import Workbook
        except Exception as ex:
            resultado_global.value = f"openpyxl não disponível: {ex}"; page.update(); return
        wb = Workbook();
        if tipo in ('Tudo','Orçamentos'):
            orcs = _fetch_all('/api/orcamentos'); ws = wb.active; ws.title = 'Orcamentos'
            if orcs: cols=list(orcs[0].keys()); ws.append(cols); [ws.append([r.get(k,'') for k in cols]) for r in orcs]
        else:
            wb.active.title='Plan1'
        if tipo in ('Tudo','Cadastros'):
            cads = _fetch_all('/api/cadastros'); ws2 = wb.create_sheet('Cadastros')
            if cads: cols=list(cads[0].keys()); ws2.append(cols); [ws2.append([r.get(k,'') for k in cols]) for r in cads]
        if tipo in ('Tudo','Pedidos'):
            peds = _fetch_all('/api/pedidos'); ws3 = wb.create_sheet('Pedidos')
            if peds: cols=list(peds[0].keys()); ws3.append(cols); [ws3.append([r.get(k,'') for k in cols]) for r in peds]
        nome = os.path.join(pasta, f"Relatorio_{tipo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
        wb.save(nome)
        resultado_global.value = f"Relatório gerado: {nome}"; page.update()

    def _on_pick_rel(d):
        if d.path:
            _exportar_por_tipo(d.path, rel_tipo.value or 'Tudo')

    _dlg_rel = ft.FilePicker(on_result=_on_pick_rel)
    page.overlay.append(_dlg_rel)
    btn_relatorio = ft.ElevatedButton('Exportar Excel', on_click=lambda e: _dlg_rel.get_directory_path(), style=pill)
    cont_relatorios = ft.Column([ft.Text('Relatórios', weight='bold'), ft.Row([rel_tipo, btn_relatorio], spacing=10)], visible=False)

    # ===================== Usuários =====================
    usr_nome = ft.TextField(label='Nome', width=320)
    usr_email = ft.TextField(label='E-mail', width=280)
    usr_setor = ft.TextField(label='Setor', width=180)
    usr_cargo = ft.TextField(label='Cargo/Função', width=220)
    usr_user = ft.TextField(label='Usuário', width=160)
    usr_senha = ft.TextField(label='Senha', password=True, can_reveal_password=True, width=180, disabled=True)
    tipo_usuario = ft.Dropdown(label='Tipo Usuário', options=[ft.dropdown.Option('Administrador'), ft.dropdown.Option('Usuário')], value='Usuário', width=180)
    usr_admin = ft.Checkbox(label='Administrador', value=False)
    perm_orc = ft.Checkbox(label='Orçamentos', value=True)
    perm_cad = ft.Checkbox(label='Cadastros', value=True)
    perm_con = ft.Checkbox(label='Contratos', value=True)
    perm_rel = ft.Checkbox(label='Relatórios', value=True)
    perm_usr = ft.Checkbox(label='Usuários', value=False)

    def _on_tipo_change(e=None):
        is_admin = (tipo_usuario.value == 'Administrador')
        usr_admin.value = is_admin
        for cb in (perm_orc, perm_cad, perm_con, perm_rel, perm_usr):
            cb.disabled = is_admin
            if is_admin:
                cb.value = True
        page.update()
    tipo_usuario.on_change = _on_tipo_change

    usuarios_table = ft.DataTable(columns=[ft.DataColumn(ft.Text('Selecionar')), ft.DataColumn(ft.Text('Usuário')), ft.DataColumn(ft.Text('Nome')), ft.DataColumn(ft.Text('E-mail')), ft.DataColumn(ft.Text('Admin'))], rows=[])

    sel_usuario_ref = {"usuario": None}

    def _refresh_users():
        try:
            r = api_get('/api/usuarios').get('rows', [])
        except Exception:
            r = []
        rows = []
        for u in r:
            def _sel(ev, uu=u):
                sel_usuario_ref['usuario'] = uu.get('usuario')
                usr_user.value = uu.get('usuario','')
                usr_nome.value = uu.get('nome','')
                usr_email.value = uu.get('email','')
                usr_setor.value = uu.get('setor','')
                usr_cargo.value = uu.get('cargo','')
                tipo_usuario.value = 'Administrador' if uu.get('is_admin') else 'Usuário'; _on_tipo_change()
                # permissões simples
                p = (uu.get('permissoes') or '')
                if p == '*' or tipo_usuario.value=='Administrador':
                    perm_orc.value=perm_cad.value=perm_con.value=perm_rel.value=perm_usr.value=True
                else:
                    s=set([x.strip() for x in p.split(',') if x.strip()])
                    perm_orc.value = 'orcamentos' in s
                    perm_cad.value = 'cadastros' in s
                    perm_con.value = 'contratos' in s
                    perm_rel.value = 'relatorios' in s
                    perm_usr.value = 'usuarios' in s
                page.update()
            btn = ft.TextButton('Selecionar', on_click=_sel)
            rows.append(ft.DataRow(cells=[ft.DataCell(btn), ft.DataCell(ft.Text(u.get('usuario',''))), ft.DataCell(ft.Text(u.get('nome',''))), ft.DataCell(ft.Text(u.get('email',''))), ft.DataCell(ft.Text('Sim' if u.get('is_admin') else 'não'))]))
        usuarios_table.rows = rows
        page.update()

    def _salvar_usuario(e):
        perms = []
        if perm_orc.value: perms.append('orcamentos')
        if perm_cad.value: perms.append('cadastros')
        if perm_con.value: perms.append('contratos')
        if perm_rel.value: perms.append('relatorios')
        if perm_usr.value: perms.append('usuarios')
        body = {
            'usuario': usr_user.value or '', 'nome': usr_nome.value or '', 'email': usr_email.value or '',
            'setor': usr_setor.value or '', 'cargo': usr_cargo.value or '', 'senha': usr_senha.value or None,
            'is_admin': bool(tipo_usuario.value=='Administrador'), 'permissoes': '*' if tipo_usuario.value=='Administrador' else ','.join(perms)
        }
        try:
            api_post('/api/usuarios', body)
            _refresh_users()
            resultado_global.value = 'Usuário salvo.'
        except Exception as ex:
            resultado_global.value = f'Erro ao salvar Usuário: {ex}'
        page.update()

    btn_user_salvar = ft.ElevatedButton('Salvar Usuário', on_click=_salvar_usuario, style=pill)
    btn_user_buscar = ft.ElevatedButton('Buscar', on_click=lambda e: _refresh_users(), style=pill)
    def _enable_user_edit():
        try:
            for w in [usr_nome, usr_email, usr_setor, usr_cargo, usr_user, usr_senha, tipo_usuario, perm_orc, perm_cad, perm_con, perm_rel, perm_usr]:
                try: w.disabled = False
                except Exception: pass
        except Exception:
            pass
        resultado_global.value = 'Edição liberada.'
        page.update()
    def _trocar_senha(e):
        try:
            api_post('/api/usuarios/change-senha', {'usuario': current_user.get('usuario'), 'senha_atual': my_old.value or '', 'senha_nova': my_new.value or '', 'force': False})
            resultado_global.value = 'Senha alterada.'
        except Exception as ex:
            resultado_global.value = f'Erro ao alterar senha: {ex}'
        page.update()
    btn_user_editar = ft.ElevatedButton('Editar Usuários', on_click=lambda e: require_admin(_enable_user_edit), style=pill)
    btn_user_editar = ft.ElevatedButton('Editar Usuários', on_click=lambda e: require_admin(_enable_user_edit), style=pill)
    
    # Campos para alteração de senha do usuário atual
    my_old = ft.TextField(label='Senha atual', password=True, can_reveal_password=True, width=180)
    my_new = ft.TextField(label='Nova senha', password=True, can_reveal_password=True, width=180)
    btn_change = ft.ElevatedButton('Alterar minha senha', on_click=_trocar_senha, style=pill)

    # Reset de senha (admin)
    adm_user = ft.TextField(label='Usuário', width=160)
    adm_new = ft.TextField(label='Nova senha', password=True, can_reveal_password=True, width=180)
    def _reset_admin(e):
        try:
            api_post('/api/usuarios/change-senha', {'usuario': adm_user.value or '', 'senha_atual': None, 'senha_nova': adm_new.value or '', 'force': True})
            resultado_global.value = 'Senha redefinida.'
        except Exception as ex:
            resultado_global.value = f'Erro ao redefinir senha: {ex}'
        page.update()
    btn_reset = ft.ElevatedButton('Redefinir senha (ADM)', on_click=_reset_admin, style=pill)

    cont_usuarios = ft.Column([
        ft.Text('Usuários', weight='bold'),
        ft.Row([usr_nome, usr_email, usr_setor, usr_cargo], wrap=True, spacing=10),
        ft.Row([usr_user, usr_senha, tipo_usuario], wrap=True, spacing=10),
        ft.Row([perm_orc, perm_cad, perm_con, perm_rel, perm_usr], wrap=True, spacing=10, alignment=ft.MainAxisAlignment.END),
        ft.Row([btn_user_salvar, btn_user_buscar, btn_user_editar], spacing=10),
        ft.Divider(),
        ft.Text('Alterar Senha', weight='bold'),
        ft.Row([my_old, my_new, btn_change], wrap=True, spacing=10),
        ft.Divider(),
        ft.Text('Reset de Senha (ADM)', weight='bold'),
        ft.Row([adm_user, adm_new, btn_reset], wrap=True, spacing=10),
        usuarios_table,
    ], visible=False)

    _refresh_users()

    page.add(nav, cont_orcamento, cont_cadastro, cont_contrato, cont_relatorios, cont_usuarios, resultado_global)

    # Correção rápida de caracteres corrompidos na interface
    _REPL = {
        'Orçamentos':'Orçamentos','Usuários':'Usuários','Usuário':'Usuário','Relatórios':'Relatórios',
        'Serviço':'Serviço','Preço':'Preço','Orçamento':'Orçamento','Impressão':'Impressão','não':'não',
        'Endereço':'Endereço','Número':'Número','Município':'Município','Período':'Período','Centímetros':'Centímetros'
    }
    def _fix_str(s):
        if isinstance(s,str) and '' in s:
            for k,v in _REPL.items():
                s=s.replace(k,v)
        return s
    def _fix_ctrl(ctrl):
        for a in ('text','label','value','hint_text','title'):
            if hasattr(ctrl,a):
                try:
                    val=getattr(ctrl,a)
                    if isinstance(val,str) and '' in val:
                        setattr(ctrl,a,_fix_str(val))
                except Exception:
                    pass
        if hasattr(ctrl,'content') and hasattr(ctrl.content,'value'):
            if isinstance(ctrl.content.value,str) and '' in ctrl.content.value:
                ctrl.content.value=_fix_str(ctrl.content.value)
        for name in ('controls','rows','columns'):
            if hasattr(ctrl,name):
                try:
                    for ch in getattr(ctrl,name) or []:
                        _fix_ctrl(ch)
                except Exception:
                    pass
    for root in page.controls:
        _fix_ctrl(root)
    page.update()

# =========================================================
#                         BOOT
# =========================================================
if __name__ == "__main__":
    init_excel()
    ft.app(target=main)